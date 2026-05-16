# -*- coding: utf-8 -*-
"""Rebuild chapters 4-6 + refs + thanks with expanded content."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = "D:/.develop/edu-achievement01/edu-achievement-master/Academic/my-thesis.docx"
doc = Document(DOCX)

def rf(run, cn, en="Times New Roman", sz=Pt(12), b=False):
    run.font.size = sz; run.font.name = en; run.bold = b
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf1 = rPr.find(qn("w:rFonts"))
    if rf1 is None:
        rf1 = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf1)
    rf1.set(qn("w:eastAsia"), cn)
    rf1.set(qn("w:ascii"), en); rf1.set(qn("w:hAnsi"), en)

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "\u9ed1\u4f53", sz=Pt(16)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0  # LEFT

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "\u9ed1\u4f53", sz=Pt(15)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h3(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "\u9ed1\u4f53", sz=Pt(14)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def body(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(t)
    rf(r, "\u5b8b\u4f53", sz=Pt(12)); p.paragraph_format.line_spacing = 1.5

def add_table(doc, headers, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    te = tbl._tbl
    tblPr = te.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); te.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement("w:"+edge)
        b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000")
        borders.append(b)
    tblPr.append(borders)
    for i,h in enumerate(headers):
        c = tbl.cell(0,i); c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        rf(r, "\u5b8b\u4f53", sz=Pt(9), b=True)
        c.paragraphs[0].paragraph_format.line_spacing = 1.5
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = tbl.cell(ri+1,ci); c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(val)
            rf(r, "\u5b8b\u4f53", sz=Pt(9))
            c.paragraphs[0].paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

# =============================================
# DATABASE TABLES (re-add + new)
# =============================================
# Find where chapter 3.3.2 ends - look for 3.4 heading
# Actually these tables go before chapter 4, so just add them now
# They will appear between chapter 3.5 and chapter 4

print("Adding DB tables...")
body(doc, "\u88683-1 \u6559\u5b66\u6210\u679c\u4e3b\u8868(edu_achievement)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["achievement_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["title","varchar(200)","\u6210\u679c\u6807\u9898"],
     ["content","text","\u6210\u679c\u5185\u5bb9(\u5bcc\u6587\u672cHTML)"],
     ["file_url","varchar(1500)","\u8bc1\u660e\u6750\u6599\u6587\u4ef6\u8def\u5f84"],
     ["teacher_id","bigint","\u6559\u5e08ID\uff0c\u5173\u8054sys_user"],
     ["college_id","bigint","\u5b66\u9662ID\uff0c\u5173\u8054sys_dept"],
     ["status","char(1)","0\u8349\u7a3f/2\u5ba1\u6838\u4e2d/3\u5df2\u901a\u8fc7/4\u5df2\u9a73\u56de"],
     ["category","char(1)","\u6210\u679c\u7c7b\u578b(1~6)"],
     ["level","varchar(30)","\u7533\u62a5\u7b49\u7ea7"],
     ["create_time","datetime","\u521b\u5efa\u65f6\u95f4"],
     ["del_flag","char(1)","\u5220\u9664\u6807\u5fd7(0\u5b58\u57282\u5220\u9664)"]])

body(doc, "\u88683-2 \u5ba1\u6838\u8bb0\u5f55\u8868(edu_audit_record)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["record_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["achievement_id","bigint","\u5173\u8054\u6210\u679cID"],
     ["audit_level","char(1)","\u5ba1\u6838\u7ea7\u522b"],
     ["audit_result","char(1)","1\u901a\u8fc7/2\u9a73\u56de"],
     ["audit_opinion","varchar(500)","\u5ba1\u6838\u610f\u89c1"],
     ["auditor_id","bigint","\u5ba1\u6838\u4ebaID"],
     ["auditor_name","varchar(64)","\u5ba1\u6838\u4eba\u59d3\u540d"],
     ["create_time","datetime","\u5ba1\u6838\u65f6\u95f4"]])

body(doc, "\u88683-3 \u95e8\u6237\u65b0\u95fb\u8868(edu_news)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["news_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["title","varchar(120)","\u65b0\u95fb\u6807\u9898"],
     ["summary","varchar(255)","\u65b0\u95fb\u6458\u8981"],
     ["cover_image","varchar(500)","\u5c01\u9762\u56feURL"],
     ["content","longtext","\u5bcc\u6587\u672c\u6b63\u6587"],
     ["publish_time","datetime","\u53d1\u5e03\u65f6\u95f4"],
     ["view_count","bigint","\u6d4f\u89c8\u91cf"],
     ["publish_portal","char(1)","\u662f\u5426\u53d1\u5e03\u5230\u95e8\u6237"],
     ["sort_weight","int","\u6392\u5e8f\u6743\u91cd"],
     ["notice_type","char(1)","1\u901a\u77e5\u516c\u544a/2\u65b0\u95fb\u52a8\u6001"]])

# NEW: sys_user, sys_role, sys_menu tables
body(doc, "\u88683-4 \u7cfb\u7edf\u7528\u6237\u8868(sys_user)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["user_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["user_name","varchar(30)","\u767b\u5f55\u8d26\u53f7"],
     ["nick_name","varchar(30)","\u7528\u6237\u6635\u79f0"],
     ["password","varchar(100)","\u5bc6\u7801(BCrypt\u52a0\u5bc6)"],
     ["phonenumber","varchar(11)","\u624b\u673a\u53f7\u7801"],
     ["email","varchar(50)","\u7535\u5b50\u90ae\u7bb1"],
     ["sex","char(1)","\u6027\u522b(0\u7537/1\u5973/2\u672a\u77e5)"],
     ["dept_id","bigint","\u90e8\u95e8ID"],
     ["status","char(1)","\u72b6\u6001(0\u6b63\u5e38/1\u505c\u7528)"],
     ["del_flag","char(1)","\u5220\u9664\u6807\u5fd7"],
     ["create_time","datetime","\u521b\u5efa\u65f6\u95f4"]])

body(doc, "\u88683-5 \u7cfb\u7edf\u89d2\u8272\u8868(sys_role)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["role_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["role_name","varchar(30)","\u89d2\u8272\u540d\u79f0"],
     ["role_key","varchar(100)","\u89d2\u8272\u6743\u9650\u5b57\u7b26\u4e32"],
     ["role_sort","int","\u663e\u793a\u987a\u5e8f"],
     ["data_scope","char(1)","\u6570\u636e\u8303\u56f4(1~5)"],
     ["status","char(1)","\u89d2\u8272\u72b6\u6001(0\u6b63\u5e38)"],
     ["create_time","datetime","\u521b\u5efa\u65f6\u95f4"]])

body(doc, "\u88683-6 \u7cfb\u7edf\u83dc\u5355\u8868(sys_menu)")
add_table(doc, ["\u5b57\u6bb5\u540d","\u7c7b\u578b","\u8bf4\u660e"],
    [["menu_id","bigint","\u4e3b\u952eID\uff0c\u81ea\u589e"],
     ["menu_name","varchar(50)","\u83dc\u5355\u540d\u79f0"],
     ["parent_id","bigint","\u7236\u83dc\u5355ID"],
     ["perms","varchar(100)","\u6743\u9650\u6807\u8bc6\u7b26"],
     ["menu_type","char(1)","\u7c7b\u578b(M\u76ee\u5f55/C\u83dc\u5355/F\u6309\u94ae)"],
     ["path","varchar(200)","\u8def\u7531\u5730\u5740"],
     ["component","varchar(255)","\u524d\u7aef\u7ec4\u4ef6\u8def\u5f84"],
     ["icon","varchar(100)","\u83dc\u5355\u56fe\u6807"],
     ["order_num","int","\u663e\u793a\u987a\u5e8f"],
     ["status","char(1)","\u83dc\u5355\u72b6\u6001(0\u663e\u793a/1\u9690\u85cf)"]])

print("DB tables done. Building chapters 4-6...")
doc.save(DOCX)
print(f"After tables: {len(doc.paragraphs)} paragraphs")
