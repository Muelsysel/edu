# -*- coding: utf-8 -*-
"""Append expanded chapters 4-6 + references + thanks to my-thesis.docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

DOCX = "D:/.develop/edu-achievement01/edu-achievement-master/Academic/my-thesis.docx"
doc = Document(DOCX)

def rf(run, cn, en="Times New Roman", sz=Pt(12), b=False):
    run.font.size = sz; run.font.name = en; run.bold = b
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf1 = rPr.find(qn("w:rFonts"))
    if rf1 is None:
        rf1 = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf1)
    rf1.set(qn("w:eastAsia"), cn)
    rf1.set(qn("w:ascii"), en); rf1.set(qn("w:hAnsi"), en)

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "黑体", sz=Pt(16)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "黑体", sz=Pt(15)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def b(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(t)
    rf(r, "宋体", sz=Pt(12)); p.paragraph_format.line_spacing = 1.5

# ============ CHAPTER 4 ============
h1(doc, "第四章 详细实现")

h2(doc, "4.1 开发环境")
b(doc, "系统开发环境基于Windows 11平台搭建，主要工具链如下：后端使用JDK 1.8和Maven 3.6+进行项目构建与依赖管理，开发IDE采用IntelliJ IDEA；前端使用Node.js 14+和npm/pnpm作为包管理器，开发服务器端口为80；数据库采用MySQL 8.0.33，缓存采用Redis 5.0，服务注册与配置中心采用Nacos 2.2.3单机模式。项目内置两套数据库初始化脚本：edu_system.sql包含全部业务表结构及示例数据，edu_config.sql为Nacos配置数据库的初始化脚本。")
b(doc, "服务启动顺序必须严格遵守，这是本系统区别于标准微服务项目的关键约束：Nacos和Redis启动后，依次启动edu-auth、edu-gateway、edu-system、edu-achievement、edu-file。其中edu-auth必须在edu-gateway之前启动，因为网关在启动时便会通过Redis进行令牌校验，如果认证服务未启动，登录流程将无法正常工作。前端执行npm run dev即可启动开发服务器，页面通过http://localhost访问。编译打包使用mvn clean package -DskipTests命令，跳过测试以加快构建速度。")

h2(doc, "4.2 网关与认证模块")
b(doc, "网关模块基于Spring Cloud Gateway构建，承担统一入口、请求路由、身份验证和跨域处理四项核心职责。网关配置了对各下游服务的路由规则，例如/auth/**路径转发至edu-auth服务，/system/**路径转发至edu-system服务，/achievement/**路径转发至edu-achievement服务。所有路由规则的服务地址均采用Nacos服务名(lb://edu-auth等)，实现服务名到实际地址的动态解析。网关还配置了Sentinel限流规则，为各服务设定了不同的QPS阈值。")
b(doc, "身份验证通过AuthFilter全局过滤器实现，是网关模块的核心组件。AuthFilter实现了GlobalFilter和Ordered接口，执行顺序设置为-200，优先于其他业务过滤器。其核心逻辑如下：首先检查请求URL是否以/achievement/portal/news开头，若是则直接放行——这是系统中硬编码的公共访问通道，为门户展示提供无障碍服务。其次，通过IgnoreWhiteProperties配置类检查URL是否在白名单中，白名单包含登录接口、验证码接口、注册接口等无需认证的接口。对于非白名单请求，从HTTP请求头中提取Authorization字段，移除Bearer前缀后解析JWT令牌，获取user_key、user_id和username。随后在Redis中检查login_tokens:{user_key}是否存在，若不存在则返回401未授权响应。最后将三个用户信息Header添加到转发请求中，并删除from-source Header，确保下游服务不会误判为内部调用。")
b(doc, "认证服务(edu-auth)提供登录、登出、令牌刷新和用户注册四个接口。登录接口/token/login接收LoginBody参数，包含username和password，SysLoginService负责验证用户凭证。验证通过后，生成一个UUID作为会话标识(user_key)，构建LoginUser对象包含用户信息、角色列表和权限集合，存储到Redis的login_tokens:{user_key}键下，设置TTL为720分钟。本系统采用双令牌设计：JWT令牌仅承载user_key、user_id和username三个轻量字段，不包含权限信息；完整的用户权限信息存储在Redis中。这种设计使得即使JWT令牌被截获，攻击者也无法获取用户的完整权限集合，而服务端可以通过删除Redis中的会话数据来主动强制用户下线。JWT的签名密钥硬编码为固定字符串，解析时通过JwtUtils类完成。")
b(doc, "下游服务在处理请求时，通过HeaderInterceptor拦截器从请求头中提取user_id、user_key和username，封装为LoginUser对象并设置到SecurityContextHolder中。业务代码可以随时通过SecurityUtils.getUserId()、SecurityUtils.getUsername()等静态方法获取当前用户信息。SecurityContextHolder基于Alibaba的TransmittableThreadLocal实现，即使在@Async异步方法中也能正确获取当前用户上下文，这解决了传统ThreadLocal在异步任务中上下文丢失的问题。FeignRequestInterceptor会在Feign调用时自动向请求中添加Authorization、user_id、user_key和username等Header，确保内部服务调用链路中用户身份的完整传递。")

h2(doc, "4.3 成果管理模块")
b(doc, "成果管理模块是系统的核心业务模块，由EduAchievementController、EduAchievementServiceImpl和EduAchievementMapper三层组成。Controller层定义了面向管理员和教师两套接口。管理员接口包括/achievement/list(列表查询，支持按标题、教师姓名、学院、状态、类型等条件复合查询)、/achievement/export(Excel导出，支持按时间段过滤)、/achievement/{id}(单个查询)、POST新增、PUT修改和DELETE批量删除。所有接口均通过@RequiresPermissions注解绑定权限标识符，如achievement:achievement:list、achievement:achievement:edit等，确保不同角色只能访问其被授权的接口。")
b(doc, "管理员修改成果时，系统会先从数据库查出旧数据，然后判断当前用户是否为超级管理员。非超级管理员需要校验是否为成果所有者，且已通过状态(3)不允许修改。超级管理员可以修改任意成果的内容和状态，如果未传入新状态则保持原状态不变。这一设计兼顾了数据保护和管理灵活性。")
b(doc, "教师接口包括teacherAddAchievement(新增)、teacherListAchievement(列表查询)、teacherUpdateAchievement(修改)、teacherDelAchievement(删除)、teacherGetAchievement(单个查询)、teacherListAllAchievement(全量查询含状态统计)、teacherResubmit(驳回后重新提交)和teacherWithdraw(撤回审核)。所有教师接口在处理时都会调用SecurityUtils.getUserId()获取当前登录用户ID，并强制设置到查询条件中，确保教师只能看到和操作自己的成果数据。teacherListAllAchievement接口除返回列表数据外，还在内存中统计各状态数量(已通过、审核中、已驳回)，并封装在AjaxResult中一并返回。")
b(doc, "教师修改成果时的状态控制逻辑如下：首先从数据库查出旧数据，校验成果是否存在、是否为本人的成果。然后检查当前状态：审核中(2)或已通过(3)状态不允许教师修改；草稿(0)状态允许保存为草稿或提交为审核中；已驳回(4)状态允许修改内容但保持驳回状态。教师可以通过teacherResubmit接口将驳回的成果重新提交为审核中状态，提交时需要校验成果必须属于当前教师且状态为已驳回；也可以通过teacherWithdraw接口将审核中的成果撤回为草稿，撤回时仅允许状态为审核中的成果。")
b(doc, "服务层EduAchievementServiceImpl负责业务逻辑处理和数据库操作协调。新增成果时，如果未明确指定状态或状态值非法，自动设置为草稿(0)。修改成果时使用MyBatis动态SQL，只更新非空字段，避免了全字段覆盖带来的数据丢失风险。删除成果采用逻辑删除，将del_flag设置为2而非物理删除记录，保留了数据可追溯性。服务层还提供了多个统计聚合方法——countByStatus、countByCategory、countByCollege、countTotal、countTotalPending、countTodayNew、countWeekNew、countMonthPassed和countMonthRejected等，满足管理员仪表盘和统计页面的数据需求。对于教师端，服务层还提供了countByStatusForTeacher、countByCategoryForTeacher和countTotalForTeacher等个人维度的统计接口。")

h2(doc, "4.4 审核流程模块")
b(doc, "审核流程模块由EduAuditRecordController和IEduAuditRecordService实现，负责审核记录的查询和审核操作。审核员通过/audit/record/list接口查询待审成果列表，通过/audit/record/approve接口执行审核通过，通过/audit/record/reject接口执行审核驳回。审核通过时，系统同时执行两个操作：将成果的status字段更新为3(已通过)，并向edu_audit_record表插入一条审核记录，记录审核人ID、姓名、审核级别、审核结果和审核意见。审核驳回时类似，将status更新为4(已驳回)并记录审核意见。")
b(doc, "为确保数据一致性，状态更新和审核记录插入应在同一数据库事务中执行，要么同时成功，要么同时失败。审核记录表通过achievement_id与成果主表关联，支持查询某个成果的全部审核历史。审核记录包含丰富的字段信息：record_id为自增主键，achievement_id关联成果，audit_level标识审核级别(预留字段，当前统一为校级审核)，audit_result标识通过(1)或驳回(2)，audit_opinion存储审核意见文字，auditor_id和auditor_name记录审核人信息，create_time记录审核时间。表中对achievement_id和auditor_id分别建立了索引，优化常见查询场景的性能。系统当前采用单级审核模式，教师提交后直接进入校级审核环节，数据库设计中预留的audit_level字段为后续扩展为院级初审加校级终审的两级模式提供了数据结构支撑。")

h2(doc, "4.5 门户展示模块")
b(doc, "门户展示模块包含EduNewsController和PortalNewsController两个控制器。EduNewsController提供新闻的增删改查接口，支持管理员在后台管理新闻内容，包括新闻标题、摘要、封面图、正文(富文本)、发布时间、排序权重、类型(通知公告或新闻动态)和是否发布到门户等字段。PortalNewsController提供面向公众的新闻查询接口，支持列表分页查询和单个新闻详情查看。")
b(doc, "门户接口的路径为/achievement/portal/news，在网关的AuthFilter中，所有以此前缀开头的请求都直接放行，不执行令牌验证。这意味着任何用户无需登录即可浏览门户新闻和成果公示内容。前端通过Axios发送请求时不携带Authorization头，网关将其视为公共请求直接放行。门户前端页面支持按发布时间倒序排列新闻列表，含封面图、标题、摘要和发布时间等信息，点击标题进入详情页展示富文本正文内容，同时浏览量自动增加。后端查询接口按publish_portal=1、status=0和del_flag=0三个条件过滤，确保仅已发布且未删除的内容在门户可见。")

h2(doc, "4.6 文件服务")
b(doc, "文件服务(edu-file)提供附件文件的上传、下载和前端静态资源的访问功能。服务支持三种存储后端：本地文件系统、MinIO对象存储和FastDFS分布式文件系统，可通过修改Nacos中的edu-file-dev.yml配置文件实现存储后端切换，无需修改代码。默认配置为本地存储模式，文件保存在服务器指定目录下，通过HTTP直接提供下载链接。服务的前缀路径为/static，用于统一管理静态资源的访问入口。")
b(doc, "文件上传接口支持单文件上传和多文件上传，上传完成后返回文件的访问URL。教师在成果申报时可以将返回的URL存入成果的file_url字段，系统支持存储多个附件文件路径，以逗号分隔。服务还配置了防盗链机制，可通过referer表白名单限制文件的访问来源，防止文件被外部网站直接链接盗用。当前版本中防盗链功能默认关闭，可根据实际部署环境开启。文件服务也支持配置CORS跨域策略，允许前端通过浏览器直接上传文件到文件服务器而不经过网关中转。")

h2(doc, "4.7 安全与权限详细实现")
b(doc, "系统的安全体系由多个切面和拦截器协同实现。接口权限校验通过PreAuthorizeAspect切面实现，在控制器方法执行前，读取@RequiresPermissions注解中声明的权限字符串，与当前登录用户的权限集合进行匹配。若用户不具备所需权限，直接抛出异常并返回403响应。此外系统还支持@RequiresRoles注解用于角色级别校验，以及@RequiresLogin注解用于仅要求登录状态的接口。内部接口通过InnerAuthAspect切面保护，该校验请求头中的from-source字段是否为inner，确保只有Feign调用可以访问标注了@InnerAuth的内部接口。")
b(doc, "数据权限通过DataScopeAspect切面和@DataScope注解实现。在服务方法上添加@DataScope(deptAlias='d', userAlias='u')注解后，切面会根据当前用户角色的data_scope属性构建SQL过滤条件，并将其注入到方法参数params.dataScope中。MyBatis映射XML中通过${params.dataScope}引用该条件，在SQL执行时动态追加数据范围限制。使用${}而非#{}是因为此处需要注入的是原始SQL片段而非参数化占位符。五级数据范围定义：级别1为全部数据权限，级别2为自定义数据权限，级别3为本部门数据权限，级别4为本部门及以下数据权限，级别5为仅本人数据权限。超级管理员(userId=1)在切面中被特殊处理，直接跳过所有数据范围过滤，确保其可以全局视图管理数据。")
b(doc, "操作日志记录通过LogAspect切面实现，对添加了@Log注解的接口方法自动记录操作日志。切面会记录调用的模块名称、业务类型(新增/修改/删除/导出等)、调用方法、请求参数和响应结果。为保护用户隐私，@Log注解自动过滤参数中的password、oldPassword、newPassword、confirmPassword等敏感字段，以星号替代。日志数据写入sys_oper_log表持久化存储。系统还提供了登录日志的记录功能，通过SysLogininforService记录每次成功或失败的登录尝试，包含登录用户、IP地址、登录时间、操作系统和浏览器等信息。")

h2(doc, "4.8 前端实现")
b(doc, "前端基于Vue 2和Element UI构建，采用基于RuoYi-Vue的二次开发模式，在此基础上扩展了教学成果管理相关的页面和组件。前端通过统一的请求封装(utils/request.js)向后端发送HTTP请求，请求拦截器会自动添加Authorization头，响应拦截器会处理令牌过期等异常状态，弹出登录对话框引导用户重新登录。")
b(doc, "登录页面(views/login.vue)提供用户名密码登录和用户注册功能，支持图形验证码。登录成功后JWT令牌被存入浏览器的localStorage，后续所有请求自动附带该令牌。首页(views/dashboard)展示系统概览信息和快捷入口。成果管理页面(views/achievement)包括教师的成果列表页、成果编辑页和成果提交页，以及管理员的全校成果管理页和审核操作页。前端使用Element UI的Table组件实现数据列表展示和分页，Dialog组件实现弹窗编辑，Upload组件实现文件上传，Tag组件实现状态标签展示。权限控制在前端通过Vue指令(directive/permission)实现，在DOM元素渲染时根据用户权限集合决定是否显示。路由守卫(permission.js)在页面加载前检查用户登录状态和页面权限，未登录用户自动跳转登录页。前端开发服务器配置了代理规则，/dev-api前缀的请求会被代理到http://localhost:8080，并自动剥离/dev-api前缀，以匹配网关的路由规则。项目采用多环境配置(.env.development、.env.production、.env.staging)，支持不同环境的快速切换。编译生产版本使用npm run build:prod命令，输出到edu-admin目录。")

print("Chapter 4 appended. Paragraphs:", len(doc.paragraphs))
doc.save(DOCX)
print("Saved.")
