# -*- coding: utf-8 -*-
"""Complete rebuild: fix headings + abstracts + expanded chapters 4-6 + refs + thanks"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCX = "D:/.develop/edu-achievement01/edu-achievement-master/Academic/my-thesis.docx"

def rf(run, cn, en="Times New Roman", sz=Pt(12), b=False):
    run.font.size = sz; run.font.name = en; run.bold = b
    run.font.color.rgb = RGBColor(0, 0, 0)
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf1 = rPr.find(qn("w:rFonts"))
    if rf1 is None:
        rf1 = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf1)
    rf1.set(qn("w:eastAsia"), cn)
    rf1.set(qn("w:ascii"), en); rf1.set(qn("w:hAnsi"), en)

def h1(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "\u9ed1\u4f53", sz=Pt(16)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h2(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t)
    rf(r, "\u9ed1\u4f53", sz=Pt(15)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def body(doc, t):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(t)
    rf(r, "\u5b8b\u4f53", sz=Pt(12)); p.paragraph_format.line_spacing = 1.5

# ============================================
# STEP 1: Fix heading alignment in existing paragraphs
# ============================================
doc = Document(DOCX)
for p in doc.paragraphs:
    t = p.text.strip()
    if not t or not p.runs: continue
    r = p.runs[0]
    if len(t) < 12 and "\u7ae0" in t and "\u7b2c" in t:
        p.alignment = 0
        r.font.size = Pt(16)
    elif t and len(t) > 3 and t[0].isdigit() and "." in t[:4] and "\u7ae0" not in t[:6]:
        dots = t.count(".")
        if dots == 1:
            r.font.size = Pt(15); p.alignment = 0
        elif dots == 2:
            r.font.size = Pt(14); p.alignment = 0

# STEP 2: Fix Chinese abstract placeholder
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if len(t) > 100 and ("\u7814\u7a76\u76ee\u7684" in t or len(t) > 300):
        # This is a very long placeholder - try to replace
        if i < 65:  # Abstract area is in the front
            p.clear()
            r = p.add_run("\u57fa\u4e8eSpring Cloud Alibaba\u5fae\u670d\u52a1\u67b6\u6784\uff0c\u8bbe\u8ba1\u5e76\u5b9e\u73b0\u4e86\u9762\u5411\u90d1\u5dde\u8f7b\u5de5\u4e1a\u5927\u5b66\u7684\u6559\u5b66\u6210\u679c\u7ba1\u7406\u7cfb\u7edf\u3002\u7cfb\u7edf\u7531\u7f51\u5173\u670d\u52a1\u3001\u8ba4\u8bc1\u4e2d\u5fc3\u3001\u7cfb\u7edf\u7ba1\u7406\u670d\u52a1\u3001\u6210\u679c\u7ba1\u7406\u670d\u52a1\u3001\u6587\u4ef6\u670d\u52a1\u548c\u5b9a\u65f6\u4efb\u52a1\u670d\u52a1\u516d\u4e2a\u5fae\u670d\u52a1\u7ec4\u6210\uff0c\u524d\u7aef\u57fa\u4e8eVue 2\u548cElement UI\u6784\u5efa\u3002\u8bba\u6587\u9996\u5148\u5206\u6790\u4e86\u9ad8\u6821\u6559\u5b66\u6210\u679c\u7ba1\u7406\u7684\u4e1a\u52a1\u9700\u6c42\uff0c\u660e\u786e\u4e86\u6559\u5e08\u3001\u5ba1\u6838\u5458\u548c\u7ba1\u7406\u5458\u4e09\u7c7b\u89d2\u8272\u7684\u529f\u80fd\u8fb9\u754c\uff1b\u968f\u540e\u4ece\u7cfb\u7edf\u67b6\u6784\u3001\u529f\u80fd\u6a21\u5757\u3001\u6570\u636e\u5e93\u3001\u5b89\u5168\u4e0e\u6743\u9650\u56db\u4e2a\u65b9\u9762\u9610\u8ff0\u4e86\u603b\u4f53\u8bbe\u8ba1\u65b9\u6848\uff1b\u63a5\u7740\u8be6\u7ec6\u63cf\u8ff0\u4e86\u7f51\u5173\u8ba4\u8bc1\u8fc7\u6ee4\u5668\u3001\u6210\u679c\u72b6\u6001\u63a7\u5236\u3001\u5ba1\u6838\u6d41\u7a0b\u4e8b\u52a1\u4fdd\u969c\u7b49\u6838\u5fc3\u673a\u5236\u7684\u5b9e\u73b0\u65b9\u5f0f\uff1b\u6700\u540e\u901a\u8fc7\u573a\u666f\u6d4b\u8bd5\u9a8c\u8bc1\u4e86\u7cfb\u7edf\u7684\u529f\u80fd\u6b63\u786e\u6027\u3002\u7cfb\u7edf\u5b9e\u73b0\u4e86\u6559\u5b66\u6210\u679c\u4ece\u5728\u7ebf\u7533\u62a5\u3001\u5ba1\u6838\u6d41\u8f6c\u5230\u95e8\u6237\u516c\u793a\u7684\u5168\u6d41\u7a0b\u6570\u5b57\u5316\u7ba1\u7406\u3002")
            rf(r, "\u5b8b\u4f53", sz=Pt(12))
            p.paragraph_format.line_spacing = 1.5
            break

doc.save(DOCX)
print("Step 1-2 done")
