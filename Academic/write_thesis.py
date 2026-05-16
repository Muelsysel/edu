# -*- coding: utf-8 -*-
"""Write thesis content to working DOCX section by section."""
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

WORKING_DOCX = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\my-thesis.docx"

def set_run_font(run, name_cn="宋体", name_en="Times New Roman", size=Pt(12), bold=False):
    run.font.size = size
    run.font.name = name_en
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name_cn)
    rFonts.set(qn("w:ascii"), name_en)
    rFonts.set(qn("w:hAnsi"), name_en)

def set_paragraph_spacing(para, line_spacing=1.5):
    pf = para.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

def add_heading1(doc, text):
    """Add a chapter heading (e.g., 第一章 引言)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, name_cn="黑体", size=Pt(16), bold=False)
    set_paragraph_spacing(para)
    para.alignment = 1  # center
    return para

def add_heading2(doc, text):
    """Add a section heading (e.g., 1.1 选题背景)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, name_cn="黑体", size=Pt(14), bold=False)
    set_paragraph_spacing(para)
    return para

def add_heading3(doc, text):
    """Add a subsection heading (e.g., 1.1.1)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, name_cn="黑体", size=Pt(12), bold=False)
    set_paragraph_spacing(para)
    return para

def add_body(doc, text):
    """Add body paragraph in 小四号宋体."""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)  # 2字符缩进
    run = para.add_run(text)
    set_run_font(run, name_cn="宋体", size=Pt(12))
    set_paragraph_spacing(para)
    return para

def clear_template_body(doc):
    """Remove template placeholder content from the body area."""
    # Find the start of body content (after English abstract / TOC area)
    body_start = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Template body area starts after "KEY WORDS" or after TOC area
        if "KEY WORDS" in text and len(text) < 30:
            body_start = i
            break
        if "目录" in text and len(text) < 10:
            body_start = i + 5  # rough skip of TOC
            break

    if body_start == -1:
        # Fallback: remove from paragraph ~50 onwards
        body_start = 40

    # Remove all paragraphs from body_start to end
    # Iterate backwards to avoid index issues
    total = len(doc.paragraphs)
    to_remove = list(range(total - 1, body_start, -1))
    for idx in to_remove:
        p = doc.paragraphs[idx]
        p._element.getparent().remove(p._element)

    print(f"Cleared {len(to_remove)} paragraphs from index {body_start} onwards.")


# === SECTIONS TO WRITE ===
# Each section: [(heading_level, text), ...]
# level: 1=chapter, 2=section, 3=subsection, 0=body

SECTIONS = {
    "1.1": [
        (1, "第一章 引言"),
        (2, "1.1 选题背景"),
        (0, "近年来，高等教育教学改革持续深化。教育部相继发布《关于深化本科教育教学改革全面提高人才培养质量的意见》等指导性文件，明确要求完善教学成果评价机制，推动教学管理工作的数字化转型。在此背景下，教学成果的申报、评审与统计已成为衡量教师教学水平、引导教改方向的重要工作环节。"),
        (0, "郑州轻工业大学现有二十余个教学单位，涵盖工学、理学、管理学、文学、艺术学等多个学科门类，每年定期组织校级教学成果奖的申报与评审工作，涉及教改项目、教材建设、课程建设、竞赛指导等多种类别。传统的管理方式以纸质材料提交与人工汇总为主，存在以下突出问题：成果信息分散于各学院教务办公室，缺乏统一的数据汇聚与查询入口；审核流转依赖纸质签字或邮件转发，过程不够透明且周期偏长；成果统计依赖手工制表，按学院、类型、等级等多维度进行汇总分析的效率较低；已通过审核的优秀成果缺少对外展示窗口，不利于教学经验的传播与示范引领。"),
        (0, "与此同时，软件架构技术的发展为教学管理信息化建设提供了新的思路。Spring Cloud Alibaba 微服务体系为构建高内聚、低耦合的分布式应用提供了成熟的技术组件：Nacos 实现服务注册发现与统一配置管理，Spring Cloud Gateway 作为 API 网关统一入口并集中执行认证与鉴权，Redis 支持令牌缓存与会话状态管理，MyBatis 提供灵活的数据库访问映射能力。采用微服务架构将系统拆分为认证中心、系统管理、成果管理、文件服务等独立模块，各服务可独立开发、部署与扩展，较好地适应教学管理业务持续演进的需求。"),
        (0, "基于上述背景，本课题设计并实现了面向郑州轻工业大学的教学成果管理系统。系统基于 Spring Cloud 微服务架构，支持教师在线申报教学成果、校级审核流转、多维度统计分析以及成果门户公示等功能，为高校教学成果的全流程数字化管理提供了一套实用的解决方案。"),
    ],
    "1.2": [
        (2, "1.2 国内外研究现状"),
    ],
    "1.3": [
        (2, "1.3 研究内容"),
    ],
    "1.4": [
        (2, "1.4 论文结构安排"),
    ],
}

# === MAIN ===
section = sys.argv[1] if len(sys.argv) > 1 else "1.1"
doc = Document(WORKING_DOCX)

if section == "1.1":
    # First section: clear template body first
    clear_template_body(doc)

if section in SECTIONS:
    for level, text in SECTIONS[section]:
        if level == 1:
            add_heading1(doc, text)
        elif level == 2:
            add_heading2(doc, text)
        elif level == 3:
            add_heading3(doc, text)
        else:
            add_body(doc, text)
    doc.save(WORKING_DOCX)
    print(f"Section {section} written to working-thesis.docx")
else:
    print(f"Section {section} not found in SECTIONS dict")
