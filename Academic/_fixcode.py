# -*- coding: utf-8 -*-
"""Phase 3: Insert code snippets into Ch4 of thesis-final.docx"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\thesis-final.docx"
doc = Document(OUTPUT)

def body_ins(parent, pos, text):
    wp = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); wp.insert(0,pPr)
    ind = OxmlElement("w:ind"); ind.set(qn("w:firstLine"),"480"); pPr.append(ind)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:line"),"360"); sp.set(qn("w:lineRule"),"auto")
    sp.set(qn("w:before"),"0"); sp.set(qn("w:after"),"0"); pPr.append(sp)
    wr = OxmlElement("w:r"); wp.append(wr)
    wrPr = OxmlElement("w:rPr"); wr.append(wrPr)
    wrSz = OxmlElement("w:sz"); wrSz.set(qn("w:val"),"24"); wrPr.append(wrSz)
    wrRf = OxmlElement("w:rFonts")
    wrRf.set(qn("w:eastAsia"),"宋体"); wrRf.set(qn("w:ascii"),"Times New Roman")
    wrPr.append(wrRf)
    wrt = OxmlElement("w:t"); wrt.text = text; wr.append(wrt)
    parent.insert(pos, wp)
    return pos + 1

def code_ins(parent, pos, text):
    wp = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); wp.insert(0,pPr)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"),"480"); pPr.append(ind)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:line"),"288"); sp.set(qn("w:lineRule"),"auto")
    sp.set(qn("w:before"),"0"); sp.set(qn("w:after"),"0"); pPr.append(sp)
    wr = OxmlElement("w:r"); wp.append(wr)
    wrPr = OxmlElement("w:rPr"); wr.append(wrPr)
    wrSz = OxmlElement("w:sz"); wrSz.set(qn("w:val"),"18"); wrPr.append(wrSz)
    wrRf = OxmlElement("w:rFonts")
    wrRf.set(qn("w:eastAsia"),"Consolas"); wrRf.set(qn("w:ascii"),"Consolas")
    wrPr.append(wrRf)
    wrt = OxmlElement("w:t"); wrt.text = text; wr.append(wrt)
    parent.insert(pos, wp)
    return pos + 1

# Find Ch4 sections
sec42_idx = sec43_idx = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "4.2" in t and "网关" in t and len(t) < 20: sec42_idx = i
    if "4.3" in t and "成果管理" in t and len(t) < 20: sec43_idx = i

print(f"4.2 at [{sec42_idx}], 4.3 at [{sec43_idx}]")

# Insert AuthFilter code after 4.2
if sec42_idx:
    p4 = doc.paragraphs[sec42_idx]._element
    parent = p4.getparent()
    # Find the 3rd body paragraph after the heading
    body_paras = [i for i, p in enumerate(doc.paragraphs) if i > sec42_idx and len(p.text.strip()) > 60]
    if len(body_paras) >= 3:
        ref = doc.paragraphs[body_paras[2]]._element
        pos = list(parent).index(ref) + 1
        pos = body_ins(parent, pos, "以下为AuthFilter核心过滤逻辑的关键代码，展示了网关层令牌验证的完整流程：")
        lines = [
            '@Override',
            'public Mono<Void> filter(ServerWebExchange exchange,',
            '                         GatewayFilterChain chain) {',
            '    String url = exchange.getRequest().getURI().getPath();',
            '    // 门户公开访问通道，直接放行',
            '    if (url.startsWith("/achievement/portal/news"))',
            '        return chain.filter(exchange);',
            '    // 白名单URL放行',
            '    if (StringUtils.matches(url, ignoreWhite.getWhites()))',
            '        return chain.filter(exchange);',
            '    // 提取并解析JWT令牌',
            '    String token = getToken(request);',
            '    Claims claims = JwtUtils.parseToken(token);',
            '    if (claims == null)',
            '        return unauthorizedResponse(exchange, "令牌无效");',
            '    // Redis验证登录状态',
            '    String userkey = JwtUtils.getUserKey(claims);',
            '    if (!redisService.hasKey("login_tokens:" + userkey))',
            '        return unauthorizedResponse(exchange, "状态已过期");',
            '    // 添加转发Header，删除from-source防伪装',
            '    addHeader(mutate, USER_KEY, userkey);',
            '    addHeader(mutate, USER_ID, JwtUtils.getUserId(claims));',
            '    removeHeader(mutate, FROM_SOURCE);',
            '    return chain.filter(exchange);',
            '}',
        ]
        for l in lines: pos = code_ins(parent, pos, l)
        print("AuthFilter code inserted")

# Insert teacherEdit code after 4.3
if sec43_idx:
    body_paras2 = [i for i, p in enumerate(doc.paragraphs) if i > sec43_idx and len(p.text.strip()) > 60]
    if len(body_paras2) >= 4:
        ref2 = doc.paragraphs[body_paras2[3]]._element
        pos2 = list(parent).index(ref2) + 1
        pos2 = body_ins(parent, pos2, "以下为教师修改成果时控制器中的状态控制核心代码，展示了身份校验和状态机控制逻辑：")
        lines2 = [
            '@PutMapping("/teacherUpdateAchievement")',
            'public AjaxResult teacherEdit(@RequestBody EduAchievement dto) {',
            '    EduAchievement oldData = service.selectById(dto.getId());',
            '    // 身份校验：仅允许修改自己的成果',
            '    if (!oldData.getTeacherId().equals(',
            '            SecurityUtils.getUserId()))',
            '        return AjaxResult.error("你无权修改他人的成果");',
            '    String oldStatus = oldData.getStatus();',
            '    // 状态校验：审核中(2)或已通过(3)不可修改',
            '    if ("2".equals(oldStatus))',
            '        return AjaxResult.error("该成果审核中，不可修改");',
            '    if ("3".equals(oldStatus))',
            '        return AjaxResult.error("已通过审核，不可修改");',
            '    // 草稿(0)可提交为审核中(2)，驳回(4)保持驳回',
            '    if ("0".equals(oldStatus))',
            '        dto.setStatus("2".equals(dto.getStatus()) ? "2" : "0");',
            '    else if ("4".equals(oldStatus)) dto.setStatus("4");',
            '    dto.setUpdateBy(SecurityUtils.getUsername());',
            '    return toAjax(service.updateById(dto));',
            '}',
        ]
        for l in lines2: pos2 = code_ins(parent, pos2, l)
        print("teacherEdit code inserted")

doc.save(OUTPUT)
print(f"Saved. {len(doc.paragraphs)} paragraphs")
