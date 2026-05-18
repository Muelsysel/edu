# -*- coding: utf-8 -*-
"""Complete thesis rebuild with correct figure placement + code snippets."""
import os, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\软件学院本科毕业设计（论文）模板 2026.docx"
OUTPUT = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\thesis-final2.docx"
FIGDIR = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\figures"

doc = Document(TEMPLATE)

# ======== HELPERS ========
def rf(run, cn, en="Times New Roman", sz=Pt(12), b=False):
    run.font.size = sz; run.font.name = en; run.bold = b
    run.font.color.rgb = RGBColor(0,0,0)
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf1 = rPr.find(qn("w:rFonts"))
    if rf1 is None:
        rf1 = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf1)
    rf1.set(qn("w:eastAsia"), cn)
    rf1.set(qn("w:ascii"), en); rf1.set(qn("w:hAnsi"), en)

def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(16)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0; p.paragraph_format.page_break_before = True

def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(15)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h3(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(14)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    rf(r, "宋体", sz=Pt(12)); p.paragraph_format.line_spacing = 1.5

def code(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Pt(24)
    r = p.add_run(text)
    rf(r, "宋体", sz=Pt(9)); p.paragraph_format.line_spacing = 1.2

def add_tbl(doc, headers, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    te = tbl._tbl
    tblPr = te.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); te.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement("w:"+edge); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000")
        borders.append(b)
    tblPr.append(borders)
    for i,h in enumerate(headers):
        c = tbl.cell(0,i); c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        rf(r, "宋体", sz=Pt(9), b=True)
        c.paragraphs[0].paragraph_format.line_spacing = 1.5
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = tbl.cell(ri+1,ci); c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(val)
            rf(r, "宋体", sz=Pt(9))
            c.paragraphs[0].paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

def table_cap(doc, text):
    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run(text); rf(r, "宋体", sz=Pt(10.5), b=True)
    p.paragraph_format.line_spacing = 1.5

def fig_cap(doc, text):
    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run(text); rf(r, "宋体", sz=Pt(10.5), b=True)
    p.paragraph_format.line_spacing = 1.5

def ref_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(text); rf(r, "宋体", sz=Pt(12))
    p.paragraph_format.line_spacing = 1.5

# ======== FILL COVER ========
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i == 16: 
        p.clear(); r = p.add_run("基于Spring Cloud的高校教学成果管理系统设计与实现")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 18 and '姓名' in t:
        p.clear(); r = p.add_run("张鹏展"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 19 and '班级' in t:
        p.clear(); r = p.add_run("软件工程22-01"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 20 and '号' in t:
        p.clear(); r = p.add_run("542201120527"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 21 and '院' in t:
        p.clear(); r = p.add_run("软件学院"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 22 and '指导教师' in t:
        p.clear(); r = p.add_run("叶志伟（高级工程师）"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 23:
        p._element.getparent().remove(p._element)

# Fix task book date
for i, p in enumerate(doc.paragraphs):
    if i == 44:
        if '期' in p.text.strip():
            p.clear(); r = p.add_run("2024年5月31日"); r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)

# ======== CLEAR BODY ========
body_start = -1
for i, p in enumerate(doc.paragraphs):
    if "KEY WORDS" in p.text.strip() and len(p.text.strip()) < 40: body_start = i; break
if body_start == -1: body_start = 60

total = len(doc.paragraphs)
for idx in range(total-1, body_start-1, -1):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
for t in list(doc.tables): t._element.getparent().remove(t._element)

# ======== ABSTRACTS ========
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i == 54 and len(t) > 200:
        p.clear(); r = p.add_run("基于Spring Cloud Alibaba微服务架构，设计并实现了面向郑州轻工业大学的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成果管理服务、文件服务和定时任务服务六个微服务组成，前端基于Vue 2和Element UI构建。论文分析了高校教学成果管理的业务需求，明确了教师、审核员和管理员三类角色的功能边界；从系统架构、功能模块、数据库、安全与权限方面阐述了总体设计方案；详细描述了网关认证过滤器、成果状态控制、审核流程事务保障等核心机制；通过场景测试验证了功能正确性。系统实现了教学成果从在线申报到审核流转再到门户公示的全流程数字化管理。")
        rf(r, "宋体", sz=Pt(12))
    if i == 56 and '关键' in t:
        p.clear(); r = p.add_run("关键词：教学成果管理；微服务架构；Spring Cloud；审核流程；Vue 2")
        rf(r, "宋体", sz=Pt(12))
    if i == 64 and len(t) > 200:
        p.clear(); r = p.add_run("Based on the Spring Cloud Alibaba microservice architecture, this thesis designs and implements a teaching achievement management system for Zhengzhou University of Light Industry. The system comprises six microservices: gateway, authentication center, system management, achievement management, file service, and scheduled task service. The thesis analyzes business requirements, identifies three user roles, elaborates the overall design covering architecture, modules, database, and security, and details the implementation of gateway authentication filter, achievement status control, and audit workflow transaction guarantee. Functional scenario testing validates system correctness. The system achieves full-process digital management from online submission to portal publication.")
        rf(r, "宋体", sz=Pt(12))
    if i == 66 and 'KEY WORDS' in t:
        p.clear(); r = p.add_run("KEY WORDS: Teaching Achievement Management; Microservices Architecture; Spring Cloud; Audit Workflow; Vue 2")
        r.font.size=Pt(12); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)

for p in list(doc.paragraphs):
    t = p.text.strip()
    if ('小四' in t or '小三' in t) and '行距' in t and len(t) < 30:
        p._element.getparent().remove(p._element)

doc.save(OUTPUT)
print("Cover+Abstracts done")

# ======== CHAPTER 1: INTRODUCTION ========
h1(doc, "第一章 引言")
doc.save(OUTPUT)

h2(doc, "1.1 选题背景")
body(doc, "近年来，高等教育教学改革持续深化。教育部相继发布《关于深化本科教育教学改革全面提高人才培养质量的意见》等指导性文件，明确要求完善教学成果评价机制，推动教学管理工作的数字化转型 [12]。在此背景下，教学成果的申报、评审与统计已成为衡量教师教学水平、引导教改方向的重要工作环节。")
body(doc, "郑州轻工业大学现有二十余个教学单位，涵盖工学、理学、管理学、文学、艺术学等多个学科门类，每年定期组织校级教学成果奖的申报与评审工作，涉及教改项目、教材建设、课程建设、竞赛指导等多种类别。传统的管理方式以纸质材料提交与人工汇总为主，存在以下突出问题：成果信息分散于各学院教务办公室，缺乏统一的数据汇聚与查询入口；审核流转依赖纸质签字或邮件转发，过程不够透明且周期偏长；成果统计依赖手工制表，按学院、类型、等级等多维度进行汇总分析的效率较低；已通过审核的优秀成果缺少对外展示窗口，不利于教学经验的传播与示范引领。")
body(doc, "与此同时，软件架构技术的发展为教学管理信息化建设提供了新的思路。Spring Cloud Alibaba 微服务体系为构建高内聚、低耦合的分布式应用提供了成熟的技术组件 [5]：Nacos 实现服务注册发现与统一配置管理 [8]，Spring Cloud Gateway 作为 API 网关统一入口并集中执行认证与鉴权 [7]，Redis 支持令牌缓存与会话状态管理 [10]，MyBatis 提供灵活的数据库访问映射能力 [9]。采用微服务架构将系统拆分为认证中心、系统管理、成果管理、文件服务等独立模块，各服务可独立开发、部署与扩展 [1]，较好地适应教学管理业务持续演进的需求。")
body(doc, "基于上述背景，本课题设计并实现了面向郑州轻工业大学的教学成果管理系统。系统基于 Spring Cloud 微服务架构，支持教师在线申报教学成果、校级审核流转、多维度统计分析以及成果门户公示等功能，为高校教学成果的全流程数字化管理提供了一套实用的解决方案。")

h2(doc, "1.2 国内外研究现状")
body(doc, "教学成果管理的信息化研究可以追溯到教育管理信息系统的早期建设阶段。2014年，James Lewis和Martin Fowler系统阐述了微服务架构风格，提出将单一应用程序划分为一组小型服务，各自围绕特定业务能力构建并通过轻量级通信机制协作 [1]。这一理念深刻影响了后续教育信息化系统的技术选型。Chris Richardson在其著作中总结了微服务架构的实现模式，包括服务拆分策略、跨服务通信、数据一致性处理等关键问题的解决方案 [2]。Sam Newman从工程实践角度出发，探讨了微服务系统的构建、测试、部署和运维方法 [3]。")
body(doc, "国内方面，高校教学管理信息化起步相对较晚但进展迅速。从2003年教育部启动精品课程建设工作开始，各高校陆续建设了不同形态的教务管理系统。早期系统采用C/S架构，侧重于学籍管理、排课选课和成绩管理等教务流程，对教学成果的专项管理覆盖较为薄弱。2010年以后，以正方教务系统和青果教务系统为代表的B/S架构平台在众多高校推广使用，增加了教研项目管理功能，但成果申报多以表单形式嵌入教务流程，缺乏独立的成果审核工作流和统计分析模块。")
body(doc, "近年来，微服务架构在教育领域的应用成为研究热点。黄建平在JWT与Redis结合的单点登录系统研究中提出了一种基于双令牌的认证机制 [4]。孙卫琴在《精通Spring Cloud Alibaba》中系统介绍了微服务架构下的服务注册、配置管理和远程调用等核心技术 [5]。王志刚等人在Spring Cloud微服务实战中，结合具体案例展示了微服务架构在复杂业务场景中的应用方法 [6]。")
body(doc, "综合以上研究现状，现有教学管理系统在教学成果管理方面仍存在以下不足：多数平台仅将成果管理作为教务系统的附属模块，未能提供完整的成果全生命周期管理闭环；对多学院环境下数据权限隔离的关注不足，缺乏按部门层级精细化控制数据访问范围的能力 [20]；部分已有研究在技术方案上仍停留在理论探讨或原型验证阶段。本课题针对郑州轻工业大学多学院协同管理的实际需求，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。")

h2(doc, "1.3 研究内容")
body(doc, "本课题以郑州轻工业大学教学成果管理的实际需求为导向，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成果管理服务、文件服务和定时任务服务六个微服务组成，配合Vue 2前端框架实现用户交互。具体研究内容包括以下四个方面：")
body(doc, "（1）教学成果申报与管理。为教师提供在线申报教学成果的完整功能，支持富文本内容编辑、附件上传、草稿保存与正式提交。教师可查看自己所有申报记录，支持按状态、类型等条件筛选。成果被驳回后可重新编辑提交，审核中可主动撤回。管理员具有全局管理权限，支持数据批量导出。")
body(doc, "（2）多级审核流程。系统实现了教学成果的状态机流转逻辑，包含草稿（0）、审核中（2）、已通过（3）和已驳回（4）四种状态。审核员可对待审成果进行通过或驳回处理，驳回时需填写审核意见。被驳回的成果允许修改后重新提交，审核中的成果允许教师主动撤回，已通过的成果不允许再次修改。")
body(doc, "（3）成果统计与分析。系统提供多维度的统计分析能力，支持按成果状态、类型、所属学院等维度汇总统计，同时提供总数、今日新增、本周新增、本月通过数和本月驳回数等实时指标，为教学管理决策提供数据支撑。")
body(doc, "（4）门户信息公示。系统内置新闻管理模块，支持发布教学动态和通知公告，并提供对外的门户展示页面。审核通过的成果可自动推送到门户网站进行公示，门户内容为公开访问，无需登录即可查看。")

h2(doc, "1.4 论文结构安排")
body(doc, "本文共分为五章。第一章为引言，介绍选题背景与意义、国内外研究现状和研究内容。第二章为需求分析，分析用户角色、功能需求和非功能需求。第三章为总体设计，阐述系统架构、模块划分、数据库设计、安全方案和技术栈选型。第四章为详细实现，逐一描述各核心模块的实现细节并展示系统运行界面。第五章为测试，通过场景测试验证系统功能正确性。最后为结束语，总结已完成的工作并展望后续改进方向。")

print("Ch1 done")
doc.save(OUTPUT)

# ======== CHAPTER 2: REQUIREMENTS ========
h1(doc, "第二章 需求分析")
h2(doc, "2.1 系统概述")
body(doc, "高校教学成果管理系统是一个面向郑州轻工业大学各教学单位的网络化服务平台，服务对象包括教师、审核员和系统管理员。系统的核心业务是对教师的教学成果申报进行全流程数字化管理，从教师填写成果内容并上传证明材料开始，经审核员的通过或驳回处理，最终形成可查询、可统计、可展示的数字化成果档案。系统涵盖了教学成果管理的全生命周期：教师在成果管理模块中可以新建成果、保存草稿、编辑内容、上传附件，也可以将完成编辑的成果提交审核；审核员通过审核模块查看待审成果的详情，填写审核意见并做出决定；审核通过的成果会自动推送至门户网站进行公示。管理员可以查看全校各学院的成果申报情况，利用统计分析功能掌握整体进度。")

h2(doc, "2.2 角色分析")
body(doc, "根据教学成果管理的实际业务流程，系统设置以下三类用户角色：")
body(doc, "（一）普通教师。教师是系统的主要使用者，负责编写和提交自己的教学成果申报。教师可以新建成果记录、保存草稿、修改内容、上传证明材料附件，并在确认无误后将成果提交审核。对于审核中状态的成果，教师可执行撤回操作将其退回草稿状态；对于被驳回的成果，教师可根据审核意见修改后重新提交。已通过审核的成果不允许再次编辑，以保证审核结果的不可篡改性。教师的所有操作被限定在本人的成果范围内，系统通过强制绑定登录用户ID的方式确保数据隔离。")
body(doc, "（二）审核员。审核员拥有查看待审成果、执行审核操作和查看审核历史记录的权限。审核员在审核界面可以查看成果的完整内容、附件材料和教师信息，然后填写审核意见并选择通过或驳回。审核过程全程留痕，每一次审核操作都会生成一条审核记录，包含审核人、审核时间、审核结果和审核意见。系统当前采用单级校级审核模式，但在数据库设计中预留了审核级别字段，可以在后续版本中扩展为院级初审和校级终审的两级审核模式。")
body(doc, "（三）超级管理员。超级管理员拥有系统的最高权限，可以查看和管理全校所有教师的成果数据，不受数据权限的限制。管理员还可以管理系统用户、角色、菜单、部门、字典等基础数据，配置权限分配和数据访问范围策略，以及查看操作日志和登录记录。系统设计中超级管理员的用户ID为固定值1，在数据权限切面中绕过所有数据范围过滤。")
body(doc, "三类角色的权限体系通过基于RBAC的权限管理模块实现。每个菜单项和接口都绑定了特定的权限标识符，用户通过角色获得对应的权限集合。后端接口通过@RequiresPermissions注解和切面编程在请求进入业务逻辑前执行权限校验 [20]。")

h2(doc, "2.3 功能需求")
h3(doc, "2.3.1 成果管理")
body(doc, "成果管理模块提供教学成果的增删改查基础操作。教师通过该模块可以新建成果记录，填写成果标题、内容详述、选择成果类型和申报等级，并上传证明材料附件。成果内容支持富文本编辑，可在详述中插入图片、表格等富文本元素。系统对成果的编辑操作实施了严格的状态控制：审核中状态（2）或已通过状态（3）的成果不允许教师修改；草稿状态（0）和已驳回状态（4）的成果允许自由编辑。教师只能查看和操作自己的成果，系统在控制器层通过SecurityUtils.getUserId()强制绑定当前登录用户ID。管理员端提供全校成果的统一查看界面，支持多条件复合查询，并可执行修改、删除和数据导出操作。")
h3(doc, "2.3.2 审核流程")
body(doc, "审核流程模块负责教学成果的状态流转和审核记录管理。成果状态设计为四种：草稿（0）、审核中（2）、已通过（3）和已驳回（4）。状态转换遵循固定规则：教师新建默认为草稿，也可直接提交为审核中；审核中经审核员处理后转为已通过或已驳回；被驳回允许修改后重新提交；审核中允许主动撤回为草稿；已通过不允许任何修改。审核员需填写审核意见后才能做出决定。每一次审核操作都会在数据库中保存一条审核记录。审核记录与成果通过achievement_id关联，支持查询完整审核轨迹。数据库设计中预留的audit_level字段为后续扩展两级审核模式提供了支撑。")
h3(doc, "2.3.3 门户新闻")
body(doc, "门户新闻模块提供教学动态和通知公告的发布与展示功能。管理员可以发布两类内容，支持富文本正文编辑、封面上传、摘要填写和排序权重设置。新闻发布后可设置是否推送到门户网站进行公开展示。门户前端提供对外的新闻列表和详情页面，无需登录即可查看。这一设计通过网关层的白名单机制实现：网关在AuthFilter中检测URL前缀，当请求路径以/achievement/portal/news开头时直接放行。")
h3(doc, "2.3.4 统计分析")
body(doc, "统计分析模块提供多维度的成果数据汇总能力。统计维度包括按成果状态、成果类型、所属学院进行汇总，同时提供总数、待审核数量、今日新增、本周新增、本月通过数和本月驳回数等实时指标。统计数据通过MyBatis映射执行SQL聚合查询获取，服务层提供了countByStatus、countByCategory、countByCollege等方法。教师端同样提供个人维度的状态统计，便于教师了解自己的成果申报进展。")
h3(doc, "2.3.5 用户权限管理")
body(doc, "用户权限管理模块提供基于RBAC模型的权限控制功能，包括用户管理、角色管理、菜单管理和部门管理四个子模块。系统在权限控制方面实现了接口级和数据级两个层级的检查：接口级通过@RequiresPermissions注解和切面编程实现；数据级通过@DataScope注解实现，在MyBatis查询时动态注入数据范围条件。数据范围分为五个级别：全部数据、自定义数据、本部门数据、本部门及以下数据和仅本人数据。")

h2(doc, "2.4 非功能需求")
body(doc, "非功能需求定义了系统在可用性、安全性、可维护性和兼容性等方面应达到的质量属性。安全性方面，系统需要对所有接口请求进行身份验证，认证采用JWT与Redis结合的双令牌机制：JWT承载用户基本身份信息，Redis存储完整的登录状态和权限信息，既保证了令牌的自包含性，又实现了服务端的会话控制 [4] [19]。可靠性方面，系统需要保证核心业务操作的数据一致性，特别是成果状态转换和审核操作需要具有事务保障，当审核员对成果执行审核时，更新成果状态和插入审核记录应保持原子性。可维护性方面，各微服务需要保持独立性，单个服务的修改和重新部署不应影响其他服务的正常运行。兼容性方面，系统需要支持主流浏览器的访问，接口遵循RESTful设计风格 [14] [13]。")

print("Ch2 done. Starting Ch3 (design diagrams)...")
doc.save(OUTPUT)

# ======== CHAPTER 3: OVERALL DESIGN (with ALL design diagrams) ========
h1(doc, "第三章 总体设计")

h2(doc, "3.1 系统架构设计")
body(doc, "本系统采用基于Spring Cloud Alibaba的微服务架构，将业务功能拆分为多个独立的微服务，每个服务围绕单一业务领域构建，服务之间通过HTTP协议和Feign客户端进行通信。系统整体架构分为四个层次：用户层、网关层、业务服务层和数据层，如图3-1所示。")
fig_cap(doc, "图3-1 系统微服务架构图")
body(doc, "用户层基于Vue 2框架构建，采用Element UI组件库实现前端页面 [11]。前端通过/dev-api前缀代理到网关服务（端口8080），网关自动剥离该前缀后转发到对应的后端服务。网关层由Spring Cloud Gateway实现，作为系统的唯一入口，承担请求路由、身份验证、越权拦截和跨域处理等职责 [7]。网关通过AuthFilter全局过滤器对每个请求进行令牌验证，合法请求会被添加用户身份Header后转发到下游业务服务。")
body(doc, "业务服务层包含认证中心（edu-auth，端口9200）、系统管理服务（edu-system，端口9201）、成果管理服务（edu-achievement，端口9205）、文件服务（edu-file，端口9300）和定时任务服务（edu-job，端口9203）五个微服务。各服务通过Nacos进行服务注册与发现，配置信息统一存储在Nacos配置中心 [8]。服务启动顺序有严格要求：认证中心必须在网关之前启动，因为网关在启动时就会通过Redis验证令牌有效性，如果认证服务未启动，登录流程将无法正常工作。")
body(doc, "数据层包括MySQL数据库、Redis缓存和文件对象存储。MySQL负责持久化存储业务数据和用户信息 [9]，Redis负责登录会话缓存和在线用户管理 [10]，文件服务支持本地存储、MinIO和FastDFS三种后端，可通过配置切换。系统还配置了监控中心（edu-monitor，端口9100），基于Spring Boot Admin实现各微服务的健康检查和运行状态监控。")
body(doc, "在安全方面，系统采用了多层防护体系。网关层的AuthFilter负责统一令牌验证和请求过滤；服务层通过AOP切面实现接口权限和数据权限的两级校验 [20]；微服务间的内部调用通过Feign客户端和InnerAuth注解保证调用合法性。这形成了网关层、服务层、数据层三位一体的安全防护体系。")

h2(doc, "3.2 功能模块划分")
body(doc, "根据需求分析，系统划分为六个功能模块，如图3-2所示。")
fig_cap(doc, "图3-2 系统功能模块划分图")
body(doc, "（1）认证与权限模块。包括用户登录、登出、令牌刷新、用户注册等功能，由edu-auth服务和网关AuthFilter共同实现。登录时用户提交用户名和密码，认证服务验证凭证后生成JWT令牌，并将完整的LoginUser对象存储在Redis中，令牌过期时间为720分钟。")
body(doc, "（2）系统管理模块。包括用户管理、角色管理、菜单管理、部门管理、字典管理、通知管理和操作日志等功能，由edu-system服务承担，提供整个系统运行所需的基础数据支持。其中菜单管理采用树形结构，支持目录、菜单和按钮三级权限粒度，为前端的动态路由渲染和按钮级别的权限控制提供数据基础。")
body(doc, "（3）成果管理模块。包括教师成果申报、成果列表查询、成果编辑、附件上传、Excel导出等功能，由edu-achievement服务承担，是系统的核心业务模块。Controller层实现了面向教师和管理员的两套接口体系，通过强制绑定用户ID的方式确保数据隔离。")
body(doc, "（4）审核流程模块。包括待审成果查询、审核通过、审核驳回、审核意见填写、审核记录查询等功能。审核操作会同时更新成果状态和插入审核记录，数据库层面通过事务注解保证操作的原子性。")
body(doc, "（5）门户展示模块。包括新闻动态和通知公告的发布编辑和门户展示功能。门户内容为公开访问，通过网关的白名单机制实现，无需登录即可浏览。")
body(doc, "（6）文件服务模块。提供附件文件的上传、下载和静态资源访问功能，支持本地存储、MinIO和FastDFS三种存储后端，可通过配置文件灵活切换。同时配置了防盗链和CORS跨域控制机制。")

h2(doc, "3.3 数据库设计")
h3(doc, "3.3.1 E-R模型")
body(doc, "系统的数据库设计围绕教学成果管理的核心业务展开，主要实体包括教师（sys_user）、学院（sys_dept）、教学成果（edu_achievement）、审核记录（edu_audit_record）和门户新闻（edu_news）。实体之间的关系如下：一个教师可以申报多个教学成果，一个学院下有多个教师，一个教学成果可以有多条审核记录。系统E-R图如图3-3所示。")
fig_cap(doc, "图3-3 数据库E-R图")

h3(doc, "3.3.2 核心表结构")
body(doc, "系统核心业务表包括教学成果主表、审核记录表和门户新闻表。系统管理相关表包括用户表、角色表、菜单表、部门表、字典数据表、字典类型表、参数配置表和操作日志表。以下为十张主要数据表的设计结构。")

table_cap(doc, "表3-1 教学成果主表(edu_achievement)")
add_tbl(doc, ["字段名","类型","说明"],
    [["achievement_id","bigint","主键ID，自增"],["title","varchar(200)","成果标题"],["content","text","成果内容(富文本HTML)"],
     ["file_url","varchar(1500)","证明材料文件路径"],["teacher_id","bigint","教师ID，关联sys_user"],
     ["college_id","bigint","学院ID，关联sys_dept"],["status","char(1)","0草稿/2审核中/3已通过/4已驳回"],
     ["category","char(1)","成果类型(1~6)"],["level","varchar(30)","申报等级"],
     ["create_time","datetime","创建时间"],["del_flag","char(1)","删除标志"]])
table_cap(doc, "表3-2 审核记录表(edu_audit_record)")
add_tbl(doc, ["字段名","类型","说明"],
    [["record_id","bigint","主键ID"],["achievement_id","bigint","关联成果ID"],
     ["audit_level","char(1)","审核级别"],["audit_result","char(1)","1通过/2驳回"],
     ["audit_opinion","varchar(500)","审核意见"],["auditor_id","bigint","审核人ID"],
     ["auditor_name","varchar(64)","审核人姓名"],["create_time","datetime","审核时间"]])
table_cap(doc, "表3-3 门户新闻表(edu_news)")
add_tbl(doc, ["字段名","类型","说明"],
    [["news_id","bigint","主键ID"],["title","varchar(120)","新闻标题"],
     ["summary","varchar(255)","新闻摘要"],["cover_image","varchar(500)","封面图URL"],
     ["content","longtext","富文本正文"],["publish_time","datetime","发布时间"],
     ["view_count","bigint","浏览量"],["publish_portal","char(1)","是否发布到门户"],
     ["sort_weight","int","排序权重"],["notice_type","char(1)","1通知公告/2新闻动态"]])
table_cap(doc, "表3-4 系统用户表(sys_user)")
add_tbl(doc, ["字段名","类型","说明"],
    [["user_id","bigint","主键ID"],["user_name","varchar(30)","登录账号"],
     ["nick_name","varchar(30)","用户昵称"],["password","varchar(100)","密码(BCrypt加密)"],
     ["dept_id","bigint","部门ID"],["email","varchar(50)","电子邮箱"],
     ["phonenumber","varchar(11)","手机号码"],["sex","char(1)","性别"],
     ["status","char(1)","状态"],["create_time","datetime","创建时间"]])
table_cap(doc, "表3-5 系统角色表(sys_role)")
add_tbl(doc, ["字段名","类型","说明"],
    [["role_id","bigint","主键ID"],["role_name","varchar(30)","角色名称"],
     ["role_key","varchar(100)","角色权限字符串"],["data_scope","char(1)","数据范围(1~5)"],
     ["role_sort","int","显示顺序"],["status","char(1)","角色状态"],["create_time","datetime","创建时间"]])
table_cap(doc, "表3-6 系统菜单表(sys_menu)")
add_tbl(doc, ["字段名","类型","说明"],
    [["menu_id","bigint","主键ID"],["menu_name","varchar(50)","菜单名称"],
     ["parent_id","bigint","父菜单ID"],["perms","varchar(100)","权限标识符"],
     ["menu_type","char(1)","M目录/C菜单/F按钮"],["path","varchar(200)","路由地址"],
     ["component","varchar(255)","前端组件路径"],["order_num","int","显示顺序"],
     ["status","char(1)","菜单状态"]])
table_cap(doc, "表3-7 部门表(sys_dept)")
add_tbl(doc, ["字段名","类型","说明"],
    [["dept_id","bigint","主键ID"],["parent_id","bigint","父部门ID"],
     ["ancestors","varchar(50)","祖级列表"],["dept_name","varchar(30)","部门名称"],
     ["order_num","int","显示顺序"],["leader","varchar(20)","负责人"],["status","char(1)","状态"]])
table_cap(doc, "表3-8 字典数据表(sys_dict_data)")
add_tbl(doc, ["字段名","类型","说明"],
    [["dict_code","bigint","主键ID"],["dict_sort","int","排序"],
     ["dict_label","varchar(100)","字典标签"],["dict_value","varchar(100)","字典键值"],
     ["dict_type","varchar(100)","字典类型"],["status","char(1)","状态"]])
table_cap(doc, "表3-9 参数配置表(sys_config)")
add_tbl(doc, ["字段名","类型","说明"],
    [["config_id","int","主键ID"],["config_name","varchar(100)","参数名称"],
     ["config_key","varchar(100)","参数键名"],["config_value","varchar(500)","参数键值"],
     ["config_type","char(1)","系统内置(Y/N)"]])
table_cap(doc, "表3-10 操作日志表(sys_oper_log)")
add_tbl(doc, ["字段名","类型","说明"],
    [["oper_id","bigint","主键ID"],["title","varchar(50)","模块标题"],
     ["business_type","int","业务类型"],["method","varchar(200)","方法名称"],
     ["oper_name","varchar(50)","操作人员"],["oper_time","datetime","操作时间"]])

h2(doc, "3.4 安全与权限设计")
body(doc, "系统的安全设计采用了多层防护体系，从网关层、服务层和数据层三个层面确保系统安全。系统用例图如图3-4所示，涵盖了教师、审核员和管理员三类角色与系统的交互关系。")
fig_cap(doc, "图3-4 系统用例图")

h3(doc, "3.4.1 认证与令牌验证流程")
body(doc, "网关层安全通过AuthFilter全局过滤器实现，该过滤器的执行顺序为-200，优先于其他业务过滤器。对于每个请求，首先检查URL是否在白名单中（包括登录接口、验证码接口和门户新闻接口），白名单中的请求直接放行；其余请求需要携带有效的JWT令牌。过滤器从HTTP请求头中提取Authorization字段，解析JWT获取用户密钥，然后在Redis中检查login_tokens:{user_key}是否存在，如果不存在则认为登录状态已过期。认证通过后，过滤器会将user_key、user_id和username三个Header添加到请求中，并移除from-source Header防止外部请求伪装成内部调用。整个认证流程如图3-5所示。")
fig_cap(doc, "图3-5 登录认证时序图")

body(doc, "认证服务（edu-auth）提供登录、登出、令牌刷新和用户注册四个接口。登录接口/token/login接收LoginBody参数，SysLoginService负责验证用户凭证。验证通过后，生成UUID作为会话标识（user_key），构建LoginUser对象包含用户信息、角色列表和权限集合，存储到Redis的login_tokens:{user_key}键下，设置TTL为720分钟。本系统采用双令牌设计：JWT令牌仅承载user_key、user_id和username三个轻量字段，不包含权限信息 [19]；完整的用户权限信息存储在Redis中。下游服务在处理请求时，通过HeaderInterceptor拦截器从请求头中提取用户信息，封装为LoginUser对象并设置到SecurityContextHolder中。SecurityContextHolder基于TransmittableThreadLocal实现，即使在@Async异步方法中也能正确获取用户上下文 [15]。")

h3(doc, "3.4.2 成果状态控制机制")
body(doc, "成果的状态流转是系统的核心业务逻辑。成果从创建到最终通过审核，需要经过严格的状态校验。教师提交修改请求时，控制器首先从数据库加载当前成果的旧数据，然后依次执行三层校验：身份校验（确认修改者为成果所有者或超级管理员）、状态校验（检查当前状态是否允许修改）、数据更新（使用MyBatis动态SQL仅更新非空字段）。状态控制的核心逻辑如图3-6所示。")
fig_cap(doc, "图3-6 成果状态控制时序图")
body(doc, "四种状态的定义和转换规则如下：草稿（0）是教师编辑的起点，在此状态下教师可以自由修改内容或直接提交为审核中。审核中（2）不允许教师修改内容，但允许主动撤回。审核员可以在审核中状态下执行通过或驳回操作。已通过（3）是终态，禁止任何修改操作，成果将自动推送至门户公示。已驳回（4）状态下教师可以查看驳回原因，根据反馈修改内容后重新提交，状态恢复为审核中。这种设计既保证了审核的严肃性，又为教师提供了修改和重新申报的机会。")

h3(doc, "3.4.3 审核事务与数据一致性")
body(doc, "审核操作涉及两个数据库表的修改：更新成果主表的status字段和向审核记录表插入新记录。为保证这两个操作的原子性，系统通过@Transactional注解将整个审核流程包裹在同一事务中执行，要么同时成功，要么同时失败。具体的执行过程如图3-7所示。")
fig_cap(doc, "图3-7 审核事务流程时序图")
body(doc, "审核员提交审核请求后，控制器首先构建审核记录对象（包含审核人ID、审核人姓名、审核结果和审核意见），然后依次调用成果服务更新状态和审核服务插入记录。两个数据库操作在MyBatis层面分别执行UPDATE和INSERT语句，由Spring事务管理器统一协调。如果任一操作失败，事务会回滚，确保不会出现状态已更新但审核记录未生成的数据不一致情况。审核记录表中对achievement_id和auditor_id分别建立了索引，优化了按成果查询审核历史和按审核人查询审核记录两个常见场景的性能。")

h3(doc, "3.4.4 数据权限与接口保护")
body(doc, "服务层安全包括接口权限和数据权限两个维度。接口权限通过@RequiresPermissions、@RequiresRoles和@RequiresLogin注解实现，由PreAuthorizeAspect切面在方法执行前进行权限校验。数据权限通过@DataScope注解实现，在MyBatis查询时动态注入数据范围条件 [20]。数据权限分为五个级别：级别1为全部数据权限，级别2为自定义数据权限，级别3为本部门数据权限，级别4为本部门及以下数据权限，级别5为仅本人数据权限。超级管理员（userId=1）在数据权限切面中被特殊处理，绕过所有数据范围过滤。微服务间的内部调用通过Feign客户端实现，FeignRequestInterceptor会自动向请求中添加认证相关Header。接口层通过@InnerAuth注解标记内部接口，InnerAuthAspect切面会校验from-source Header是否为inner，确保只有内部服务可以调用这些接口。")
body(doc, "操作日志记录通过LogAspect切面实现，对添加了@Log注解的接口方法自动记录操作日志，包括调用的模块名称、业务类型、调用方法和请求参数。为保护用户隐私，@Log注解自动过滤参数中的password、oldPassword、newPassword、confirmPassword等敏感字段。核心权限控制类的静态关系如图3-8所示。")
fig_cap(doc, "图3-8 安全与权限控制核心类图")

h2(doc, "3.5 技术栈选型")
body(doc, "本系统技术栈选型以成熟稳定和与微服务架构的配套程度为主要考量。后端基于Java 1.8和Spring Boot 2.7.18构建，使用Spring Cloud 2021.0.9和Spring Cloud Alibaba 2021.0.6.1作为微服务基础设施 [5] [6]。Nacos 2.2.3提供服务注册发现和配置管理能力 [8]。数据持久化层采用MySQL 8.0作为主数据库，MyBatis作为ORM框架，PageHelper实现分页查询，Druid作为数据库连接池 [9]。缓存层采用Redis 5+存储登录会话和业务缓存 [10]。前端基于Vue 2框架和Element UI组件库构建，采用基于RuoYi-Vue的二次开发模式，在其基础上扩展了教学成果管理相关的页面和组件 [11]。")
body(doc, "其他关键组件包括：Alibaba TransmittableThreadLocal解决异步任务中上下文传递问题 [15]；Quartz实现定时任务调度；SpringDoc OpenAPI自动生成接口文档；Apache POI处理Excel数据导出；MinIO SDK实现对象存储文件操作。整体技术栈完全基于开源组件，没有商业许可依赖。")

print("Ch3 done. Starting Ch4 (implementation + screenshots + code)...")
doc.save(OUTPUT)

# ======== CHAPTER 4: IMPLEMENTATION (screenshots + code) ========
h1(doc, "第四章 详细实现")

h2(doc, "4.1 开发环境")
body(doc, "系统开发环境基于Windows 11平台搭建。后端使用JDK 1.8和Maven 3.6+进行项目构建与依赖管理，开发IDE采用IntelliJ IDEA；前端使用Node.js 14+和npm/pnpm作为包管理器，开发服务器端口为80；数据库采用MySQL 8.0.33，缓存采用Redis 5.0，服务注册与配置中心采用Nacos 2.2.3单机模式。项目内置两套数据库初始化脚本：edu_system.sql包含全部业务表结构及示例数据，edu_config.sql为Nacos配置数据库的初始化脚本。服务启动顺序必须严格遵守：Nacos和Redis启动后，依次启动edu-auth、edu-gateway、edu-system、edu-achievement、edu-file。编译打包使用mvn clean package -DskipTests命令。")

h2(doc, "4.2 网关与认证模块")
body(doc, "网关模块基于Spring Cloud Gateway构建，承担统一入口、请求路由、身份验证和跨域处理四项核心职责。网关配置了对各下游服务的路由规则，所有路由规则的服务地址均采用Nacos服务名实现动态解析。")
body(doc, "身份验证通过AuthFilter全局过滤器实现，是网关模块的核心组件。AuthFilter实现了GlobalFilter和Ordered接口，执行顺序设置为-200。其核心逻辑如下：首先检查URL是否以/achievement/portal/news开头，若是则直接放行（门户公共通道）。其次通过IgnoreWhiteProperties检查URL是否在白名单中。对于非白名单请求，从HTTP请求头中提取Authorization字段，移除Bearer前缀后解析JWT令牌，获取user_key、user_id和username。随后在Redis中检查login_tokens:{user_key}是否存在，若不存在则返回401。最后将用户信息Header添加到转发请求中，并删除from-source Header防止外部请求伪装成内部调用。")
body(doc, "以下为AuthFilter核心过滤逻辑的关键代码：")
code(doc, """@Override
public Mono<Void> filter(ServerWebExchange exchange, GatewayFilterChain chain) {
    ServerHttpRequest request = exchange.getRequest();
    String url = request.getURI().getPath();
    // 门户公开访问通道
    if (url.startsWith("/achievement/portal/news")) {
        return chain.filter(exchange);
    }
    // 白名单放行
    if (StringUtils.matches(url, ignoreWhite.getWhites())) {
        return chain.filter(exchange);
    }
    // 提取并验证JWT
    String token = getToken(request);
    Claims claims = JwtUtils.parseToken(token);
    if (claims == null) {
        return unauthorizedResponse(exchange, "令牌无效或已过期");
    }
    // Redis验证登录状态
    String userkey = JwtUtils.getUserKey(claims);
    if (!redisService.hasKey("login_tokens:" + userkey)) {
        return unauthorizedResponse(exchange, "登录状态已过期");
    }
    // 添加转发Header
    addHeader(mutate, SecurityConstants.USER_KEY, userkey);
    addHeader(mutate, SecurityConstants.DETAILS_USER_ID, JwtUtils.getUserId(claims));
    addHeader(mutate, SecurityConstants.DETAILS_USERNAME, JwtUtils.getUserName(claims));
    // 删除from-source防止外部伪装
    removeHeader(mutate, SecurityConstants.FROM_SOURCE);
    return chain.filter(exchange);
}""")
body(doc, "上述代码展示了AuthFilter的核心工作流程：首先通过URL前缀和正则匹配实现白名单放行，然后提取JWT令牌并解析其中的用户信息，最后通过Redis验证登录状态并设置下游服务所需的请求头。系统的登录页面如图4-1所示。")
fig_cap(doc, "图4-1 系统登录页面")

h2(doc, "4.3 成果管理模块")
body(doc, "成果管理模块是系统的核心业务模块，由EduAchievementController、EduAchievementServiceImpl和EduAchievementMapper三层组成。Controller层定义了面向管理员和教师的两套接口。教师接口在控制器层通过SecurityUtils.getUserId()获取当前登录用户ID并强制绑定到查询条件，从业务逻辑层确保每个教师只能访问自己的数据。")
body(doc, "教师修改成果时的状态控制逻辑是模块的核心：首先从数据库查出旧数据，校验成果是否存在以及是否为本人所有；然后检查当前状态是否允许编辑；仅当所有校验通过后才执行数据库更新。以下为控制器中状态控制的关键代码：")
code(doc, """@PutMapping("/teacherUpdateAchievement")
public AjaxResult teacherEdit(@RequestBody EduAchievement dto) {
    // 1. 从数据库查出旧数据
    EduAchievement oldData = service.selectById(dto.getAchievementId());
    if (oldData == null) return AjaxResult.error("成果不存在");
    // 2. 身份校验：仅允许修改自己的成果
    if (!oldData.getTeacherId().equals(SecurityUtils.getUserId())) {
        return AjaxResult.error("你无权修改他人的成果");
    }
    // 3. 状态校验：审核中(2)或已通过(3)不可修改
    String oldStatus = oldData.getStatus();
    if ("2".equals(oldStatus)) return AjaxResult.error("该成果审核中，不可修改");
    if ("3".equals(oldStatus)) return AjaxResult.error("已通过审核，不可修改");
    // 4. 草稿(0)可提交为审核中(2)，驳回(4)保持驳回
    if ("0".equals(oldStatus)) dto.setStatus("2".equals(dto.getStatus()) ? "2" : "0");
    else if ("4".equals(oldStatus)) dto.setStatus("4");
    // 5. 执行更新（MyBatis动态SQL，仅更新非空字段）
    dto.setUpdateBy(SecurityUtils.getUsername());
    return toAjax(service.updateById(dto));
}""")
body(doc, "服务层EduAchievementServiceImpl负责业务逻辑处理和数据库操作协调。新增成果时如果状态非法则自动设置为草稿；修改时使用MyBatis动态SQL只更新非空字段；删除采用逻辑删除保留数据可追溯性。教师端成果管理列表页面如图4-2所示，成果申报编辑页面如图4-3所示。")
fig_cap(doc, "图4-2 教师成果列表页面")
fig_cap(doc, "图4-3 成果申报编辑页面")

h2(doc, "4.4 审核流程模块")
body(doc, "审核流程模块由EduAuditRecordController和IEduAuditRecordService实现。审核员通过/audit/record/list接口查询待审成果列表（筛选status=2），通过/audit/record/approve接口执行审核通过，通过/audit/record/reject接口执行审核驳回。审核通过时，系统同时执行两个操作：将status更新为3（已通过），并向edu_audit_record表插入一条审核记录。为确保数据一致性，两个操作包裹在同一数据库事务中。以下为审核操作的关键代码：")
code(doc, """@Transactional
public int approve(Long achievementId, EduAuditRecord record) {
    // 1. 更新成果状态为已通过(3)
    EduAchievement update = new EduAchievement();
    update.setAchievementId(achievementId);
    update.setStatus("3");
    achievementService.updateById(update);
    
    // 2. 插入审核记录
    record.setAchievementId(achievementId);
    record.setAuditLevel("2");
    record.setAuditResult("1");  // 1=通过, 2=驳回
    record.setCreateTime(new Date());
    return auditRecordMapper.insert(record);
}""")
body(doc, "审核记录表通过achievement_id与成果主表关联，支持查询某个成果的全部审核历史。每条审核记录包含审核人ID、审核人姓名、审核级别、审核结果、审核意见和审核时间等字段。审核操作界面如图4-4所示。")
fig_cap(doc, "图4-4 审核操作页面")

h2(doc, "4.5 门户展示模块")
body(doc, "门户展示模块包含EduNewsController和PortalNewsController两个控制器。EduNewsController提供新闻的增删改查接口，支持富文本正文、封面图、排序权重等完整字段管理。PortalNewsController提供面向公众的新闻查询接口，支持列表分页和详情查看。门户接口路径/achievement/portal/news在网关AuthFilter中被直接放行，无需登录即可访问。前端通过Axios发送不带Authorization头的请求，网关将其视为公共请求。门户页面支持按发布时间倒序排列新闻列表，含封面图、标题和摘要信息，点击标题进入详情页展示富文本正文，同时浏览量自动增加。门户新闻展示页面如图4-5所示。")
fig_cap(doc, "图4-5 门户新闻展示页面")

h2(doc, "4.6 安全与权限实现")
body(doc, "接口权限校验通过PreAuthorizeAspect切面实现。在控制器方法执行前，切面读取@RequiresPermissions注解中声明的权限字符串，与当前登录用户的权限集合进行匹配。若用户不具备所需权限，则抛出异常并返回403响应。系统还支持@RequiresRoles注解用于角色级别校验，以及@RequiresLogin注解用于仅要求登录状态的接口。")
body(doc, "数据权限通过DataScopeAspect切面和@DataScope注解实现。在服务方法上添加@DataScope注解后，切面会根据当前用户角色的data_scope属性构建SQL过滤条件并注入到方法参数params.dataScope中。MyBatis映射XML中通过${params.dataScope}引用该条件实现动态数据过滤。以下为DataScope关键实现代码：")
code(doc, """// DataScopeAspect 切面核心逻辑
@Around("@annotation(dataScope)")
public Object doBefore(ProceedingJoinPoint point, DataScope dataScope) {
    LoginUser loginUser = SecurityContextHolder.getLoginUser();
    if (loginUser.getUserId() == 1L) {
        // 超级管理员绕过所有数据权限过滤
        return point.proceed();
    }
    StringBuilder sqlFilter = new StringBuilder();
    switch (loginUser.getDataScope()) {
        case "1":  // 全部数据权限
            return point.proceed();
        case "2":  // 自定义数据权限
            sqlFilter.append(" AND dept_id IN (SELECT dept_id FROM sys_role_dept...)");
            break;
        case "3":  // 本部门数据权限
            sqlFilter.append(" AND dept_id = ").append(loginUser.getDeptId());
            break;
        case "4":  // 本部门及以下
            sqlFilter.append(" AND dept_id IN (SELECT dept_id FROM sys_dept WHERE ancestors LIKE...)");
            break;
        case "5":  // 仅本人数据
            sqlFilter.append(" AND create_by = '").append(loginUser.getUsername()).append("'");
            break;
    }
    // 将SQL条件注入方法参数
    setDataScopeParams(point, sqlFilter.toString());
    return point.proceed();
}""")
body(doc, "管理员端的成果管理界面如图4-6所示，支持全校成果的统一查看、多条件筛选、编辑修改、删除和Excel导出等功能。统计仪表盘如图4-7所示，展示成果总数、待审核数量等关键指标以及按状态、类型、学院的多维度统计图表。")
fig_cap(doc, "图4-6 管理员成果管理页面")
fig_cap(doc, "图4-7 统计仪表盘页面")

h2(doc, "4.7 前端实现")
body(doc, "前端基于Vue 2和Element UI构建，采用基于RuoYi-Vue的二次开发模式。前端通过统一的请求封装（utils/request.js）向后端发送HTTP请求，请求拦截器自动添加Authorization头，响应拦截器处理令牌过期等异常状态。登录页面提供用户名密码登录和用户注册功能，登录成功后JWT令牌被存入localStorage。成果管理页面使用Element UI的Table组件实现数据列表展示和分页，Dialog组件实现弹窗编辑，Upload组件实现文件上传，Tag组件实现状态标签展示。权限控制在前端通过Vue指令（directive/permission）实现，路由守卫在页面加载前检查用户登录状态和页面权限。前端开发服务器配置了代理规则，/dev-api前缀的请求被代理到http://localhost:8080并自动剥离前缀，以匹配网关的路由规则。项目采用多环境配置，支持不同环境的快速切换。")

print("Ch4 done. Starting Ch5...")
doc.save(OUTPUT)

# ======== CHAPTER 5: TESTING ========
h1(doc, "第五章 测试")

h2(doc, "5.1 测试环境")
body(doc, "系统测试在开发环境中进行，硬件环境为Windows 11桌面主机，配备32GB内存和1TB SSD存储。数据库采用MySQL 8.0.33，缓存采用Redis 5.0，注册中心采用Nacos 2.2.3单机模式，JDK版本为1.8.0。测试时所有微服务均在同一台机器上以默认端口运行，前端通过http://localhost访问。测试数据的准备通过执行edu_system.sql初始化脚本完成，该脚本包含了完整的业务表结构和多组示例数据，涵盖了全校多个学院、多种类型和多种状态的成果记录，为功能测试提供了充分的数据基础。测试采用以功能验证为主的方法，从用户角度对系统的核心业务流程进行逐项验证。")

h2(doc, "5.2 功能测试")
body(doc, "本节以场景测试的形式对系统各核心功能模块进行验证，共设计了十一组测试用例，涵盖成果申报与管理、审核流程控制、权限校验、门户展示和数据统计等关键功能。以下逐项描述测试过程与验证结果。")
body(doc, "测试用例1：教师新建成果并提交审核。前置条件为教师已登录系统，身份为普通教师角色。测试步骤：进入成果管理页面，点击新建按钮，填写成果标题、选择成果类型和申报等级，在富文本编辑器中编写成果内容，上传PDF格式的证明材料附件，点击提交按钮。预期结果：系统提示提交成功，成果列表中显示新增的成果记录，状态为审核中。实际结果与预期一致，成果成功提交并进入审核状态。")
body(doc, "测试用例2：教师编辑草稿成果。前置条件为存在一个草稿状态（0）的成果，当前登录用户为该成果的申报教师。测试步骤：在成果列表中点击编辑按钮，修改成果标题和内容，保存。预期结果：成果信息更新成功，列表显示更新后的标题。实际结果与预期一致，草稿成果可正常编辑。")
body(doc, "测试用例3：审核员审核通过成果。前置条件为存在一个审核中状态（2）的成果，审核员已登录。测试步骤：在审核页面查看成果详情和附件材料，填写审核意见后点击通过按钮。预期结果：成果状态变为已通过（3），审核记录中新增一条通过记录。实际结果与预期一致。")
body(doc, "测试用例4：审核员驳回成果。前置条件为存在一个审核中状态（2）的成果。测试步骤：填写审核意见后点击驳回按钮。预期结果：成果状态变为已驳回（4），审核记录中新增一条驳回记录，教师可查看审核意见。实际结果与预期一致。")
body(doc, "测试用例5：教师重新提交被驳回的成果。前置条件为存在一个已驳回状态（4）的成果。测试步骤：查看审核意见，根据意见修改成果内容，点击重新提交按钮。预期结果：成果状态恢复为审核中（2）。实际结果与预期一致。")
body(doc, "测试用例6：教师撤回审核中的成果。前置条件为存在一个审核中状态（2）的成果。测试步骤：在成果列表中点击撤回按钮。预期结果：成果状态变为草稿（0）。实际结果与预期一致。")
body(doc, "测试用例7：已通过成果禁止修改。前置条件为存在一个已通过状态（3）的成果。测试步骤：尝试编辑该成果。预期结果：系统拒绝操作，提示\"该成果已通过审核，不可修改\"。实际结果与预期一致。")
body(doc, "测试用例8：教师越权修改他人成果。前置条件为当前登录教师非成果所有者。测试步骤：尝试通过接口修改他人的成果。预期结果：系统拒绝操作，提示\"你无权修改他人的成果\"。实际结果与预期一致，权限校验有效。")
body(doc, "测试用例9：门户新闻公开访问。前置条件为用户未登录。测试步骤：在浏览器中直接访问门户新闻列表页和详情页。预期结果：页面正常显示，无需登录，浏览量自动增加。实际结果与预期一致。")
body(doc, "测试用例10：管理员查看统计数据。前置条件为管理员已登录。测试步骤：在首页查看关键指标，分别进入按状态、类型、学院的统计页面。预期结果：各维度数据正确显示，与数据库实际数据一致。实际结果与预期一致。")
body(doc, "测试用例11：Excel导出功能。前置条件为管理员已登录。测试步骤：设置筛选条件后点击导出按钮。预期结果：下载Excel文件，包含符合筛选条件的成果数据。实际结果与预期一致。")

h2(doc, "5.3 测试结果")
body(doc, "通过上述十一组场景测试，系统各核心功能模块均通过了功能验证。成果的增删改查操作正常，状态流转逻辑严谨，权限校验和数据隔离机制有效。在测试过程中，系统对非法操作均能正确拦截并给出提示，包括越权修改他人成果、修改已通过成果、非法状态转换等场景。审核记录完整可追溯，在审核记录列表中可以查看任一成果的完整审核历史，包括通过和驳回的全部记录。门户新闻公开访问正常，无需登录即可浏览。统计数据计算准确，各维度相互一致。Excel导出功能运行正常，导出数据与列表中显示一致。")
body(doc, "综合来看，系统的核心业务功能已达到设计预期，状态流转逻辑严谨、权限控制有效、数据一致性良好，可以支持教学成果管理的日常业务需求。系统中预留的AI查重接口（返回501状态码）和审核级别扩展字段（audit_level）为后续功能优化提供了基础。为进一步提升系统的完整度，建议在后续版本中补充性能压力测试、补充AI相似度检测能力、优化移动端展示效果，以及引入容器化部署提高运维便利性 [16] [17] [18]。")

print("Ch5 done. Writing ending + refs + thanks...")
doc.save(OUTPUT)

# ======== ENDING ========
h1(doc, "结束语")
body(doc, "本课题以郑州轻工业大学教学成果管理的实际需求为背景，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成果管理服务、文件服务和定时任务服务六个微服务组成，前端基于Vue 2和Element UI构建，数据层采用MySQL、Redis和文件对象存储。本文完成的主要工作包括：")
body(doc, "（1）分析了高校教学成果管理的业务需求，明确了教师、审核员和管理员三类角色的功能范围和权限边界，划分了成果管理、审核流程、门户新闻、统计分析和用户权限管理五个核心功能模块。")
body(doc, "（2）设计了基于微服务架构的系统整体方案，包括网关层、业务服务层和数据层的分层架构。完成了数据库E-R模型设计和十张核心业务表的结构定义，并制定了JWT+Redis双令牌认证、RBAC权限控制和五级数据范围过滤等安全方案。")
body(doc, "（3）实现了系统的各个核心模块，重点阐述了网关认证过滤器的工作原理、成果管理控制器的状态控制逻辑、审核流程的数据一致性保障以及前端页面的实现方式。通过在文中展示核心代码片段和系统运行界面，直观呈现了系统的工程实现细节。通过十一组场景测试验证了系统各模块的正确性和可用性。")
body(doc, "本系统实现了教学成果从在线申报、审核流转到门户公示的全流程数字化管理，可以为高校教学成果的统一管理和数据分析提供有效支持。系统也存在一些可以改进的方面：当前审核流程为单级校级审核，后续可扩展为院级初审加校级终审的两级模式；AI查重接口已预留但尚未实现，可对接外部查重服务引入智能分析能力；前端界面可进一步适配移动端展示；微服务运维可引入分布式链路追踪和容器化部署等方案。")

# ======== REFERENCES ========
h1(doc, "参考文献")
refs = [
    "[1] James Lewis, Martin Fowler. Microservices: a definition of this new architectural term[EB/OL]. (2014-03-25). https://martinfowler.com/articles/microservices.html.",
    "[2] Chris Richardson. Microservices Patterns: With examples in Java[M]. Shelter Island: Manning Publications, 2018: 1-50.",
    "[3] Sam Newman. Building Microservices: Designing Fine-Grained Systems[M]. Sebastopol: O'Reilly Media, 2015: 1-80.",
    "[4] 黄建平. 基于JWT与Redis的单点登录系统设计与实现[J]. 计算机技术与发展, 2020, 30(3): 25-31.",
    "[5] 孙卫琴. 精通Spring Cloud Alibaba——微服务架构与实战[M]. 北京: 北京大学出版社, 2020: 1-120.",
    "[6] 王志刚, 孙晓华. Spring Cloud微服务实战[M]. 北京: 人民邮电出版社, 2019: 45-102.",
    "[7] Spring Cloud Gateway Documentation[EB/OL]. https://docs.spring.io/spring-cloud-gateway/docs/current/reference/html/.",
    "[8] 崔永志, 贾喜平. Nacos实战与源码分析——微服务注册与配置管理[M]. 北京: 电子工业出版社, 2021: 23-89.",
    "[9] 刘斌, 陈建峰. MyBatis从入门到精通[M]. 北京: 清华大学出版社, 2017: 78-135.",
    "[10] Redis官方文档[EB/OL]. https://redis.io/docs/latest/.",
    "[11] Vue.js v2官方指南[EB/OL]. https://v2.vuejs.org/v2/guide/.",
    "[12] 教育部. 关于深化本科教育教学改革全面提高人才培养质量的意见[EB/OL]. (2019-10-08). http://www.moe.gov.cn/srcsite/A08/s7056/201910/t20191011_402759.html.",
    "[13] Roy Thomas Fielding. Architectural Styles and the Design of Network-based Software Architectures[D]. Irvine: University of California, 2000.",
    "[14] Cesare Pautasso, Olaf Zimmermann, Frank Leymann. RESTful Web Services vs. Big Web Services: Making the Right Architectural Decision[C]. Beijing: WWW 2008, 2008: 805-814.",
    "[15] 耿祥义, 张跃平. Java并发编程实战[M]. 北京: 机械工业出版社, 2015: 120-156.",
    "[16] Irakli Nadareishvili, Ronnie Mitra, Matt McLarty, Mike Amundsen. Microservice Architecture: Aligning Principles, Practices, and Culture[M]. Sebastopol: O'Reilly Media, 2016: 1-60.",
    "[17] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Boston: Addison-Wesley, 2002: 1-100.",
    "[18] Eric Evans. Domain-Driven Design: Tackling Complexity in the Heart of Software[M]. Boston: Addison-Wesley, 2003: 1-80.",
    "[19] M. Jones, J. Bradley, N. Sakimura. JSON Web Token (JWT): RFC 7519[S]. IETF, 2015. https://tools.ietf.org/html/rfc7519.",
    "[20] 李俊杰. 基于RBAC模型的权限管理系统设计与实现[D]. 成都: 电子科技大学, 2018.",
]
for rt in refs: ref_para(doc, rt)

h1(doc, "致谢")
body(doc, "在本文完成之际，首先向我的指导教师致以最诚挚的感谢。在毕业设计的整个过程中，指导教师从选题确定、技术方案论证、系统设计到论文撰写的每个环节都给予了细致的指导和建设性意见，帮助我克服了诸多技术难题和思路困惑。感谢软件学院各位任课教师在四年本科学习中的辛勤教导，为我奠定了软件开发的理论基础和实践能力。感谢实验室同学们在项目开发过程中的互助合作，大家的讨论和交流让我对微服务架构有了更深入的理解。感谢郑州轻工业大学提供了优越的学习环境和实验条件，让我能够将所学知识应用于实际项目中。最后感谢家人在求学期间的理解和鼓励，他们的无私支持是我完成学业的重要保障。")

# ======== SAVE ========
doc.save(OUTPUT)
print(f"Body complete. {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# ======== INSERT IMAGES ========
doc = Document(OUTPUT)

# Build PNG map
png_map = {}
for fn in os.listdir(FIGDIR):
    if fn.endswith('.png'):
        m = re.match(r'图(\d+)-(\d+)', fn)
        if m: png_map[f'{m.group(1)}-{m.group(2)}'] = os.path.join(FIGDIR, fn)

# Find figure captions and insert images above them
fig_captions = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'图(\d+)-(\d+)', t)
    if m and p.runs and p.runs[0].bold and len(t) < 30:
        fig_captions.append((i, m.group(1), m.group(2)))

print(f'Found {len(fig_captions)} figure captions')

inserted = 0
for para_idx, ch, fnum in reversed(fig_captions):
    key = f'{ch}-{fnum}'
    if key not in png_map: continue
    
    cap_elem = doc.paragraphs[para_idx]._element
    parent = cap_elem.getparent()
    cap_pos = list(parent).index(cap_elem)
    
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Cm(0)
    img_para.paragraph_format.space_after = Cm(0)
    
    try:
        img_para.add_run().add_picture(png_map[key], width=Inches(4.8))
        img_elem = img_para._element
        img_elem.getparent().remove(img_elem)
        parent.insert(cap_pos, img_elem)
        inserted += 1
    except: pass

doc.save(OUTPUT)
print(f'Images inserted: {inserted}')
print(f'COMPLETE: {OUTPUT}')
