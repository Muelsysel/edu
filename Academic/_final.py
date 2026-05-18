# -*- coding: utf-8 -*-
"""Final rebuild: template copy + all content + images in one pass."""
import os, sys, glob, re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATE = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\软件学院本科毕业设计（论文）模板 2026.docx"
OUTPUT = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\thesis-final.docx"
FIGDIR = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\figures"

# Open template
doc = Document(TEMPLATE)

# ======== HELPER FUNCTIONS ========
def rf(run, cn, en="Times New Roman", sz=Pt(12), b=False):
    run.font.size = sz; run.font.name = en; run.bold = b
    run.font.color.rgb = RGBColor(0,0,0)
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = run._element.makeelement(qn("w:rPr"), {})
        run._element.insert(0, rPr)
    rf1 = rPr.find(qn("w:rFonts"))
    if rf1 is None:
        rf1 = rPr.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rf1)
    rf1.set(qn("w:eastAsia"), cn)
    rf1.set(qn("w:ascii"), en); rf1.set(qn("w:hAnsi"), en)

def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(16)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(15)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def h3(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    rf(r, "黑体", sz=Pt(14)); p.paragraph_format.line_spacing = 1.5
    p.alignment = 0

def body(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(text)
    rf(r, "宋体", sz=Pt(12)); p.paragraph_format.line_spacing = 1.5

def add_tbl(doc, headers, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
    te = tbl._tbl
    tblPr = te.find(qn("w:tblPr"))
    if tblPr is None: tblPr = OxmlElement("w:tblPr"); te.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement("w:"+edge); b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
        b.set(qn("w:space"),"0"); b.set(qn("w:color"),"000000")
        borders.append(b)
    tblPr.append(borders)
    for i,h in enumerate(headers):
        c = tbl.cell(0,i); c.paragraphs[0].clear()
        r = c.paragraphs[0].add_run(h)
        rf(r, "宋体", sz=Pt(9), b=True)
        c.paragraphs[0].paragraph_format.line_spacing = 1.5
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = tbl.cell(ri+1,ci); c.paragraphs[0].clear()
            r = c.paragraphs[0].add_run(val)
            rf(r, "宋体", sz=Pt(9))
            c.paragraphs[0].paragraph_format.line_spacing = 1.5
    doc.add_paragraph()

def table_cap(doc, text):
    p = doc.add_paragraph(); p.alignment = 1
    r = p.add_run(text); rf(r, "宋体", sz=Pt(10.5), b=True)
    p.paragraph_format.line_spacing = 1.5

def ref_para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(text); rf(r, "宋体", sz=Pt(12))
    p.paragraph_format.line_spacing = 1.5

# ======== FILL COVER INFO ========
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i == 16 and '题   目' in t:
        p.clear(); r = p.add_run("基于Spring Cloud的高校教学成果管理系统设计与实现")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 18 and '学生姓名' in t:
        p.clear(); r = p.add_run("张鹏展")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 19 and '专业班级' in t:
        p.clear(); r = p.add_run("软件工程22-01")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 20:
        if '学   号' in t:
            p.clear(); r = p.add_run("542201120527")
            r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 21 and '学   院' in t:
        p.clear(); r = p.add_run("软件学院")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 22 and '指导教师' in t:
        p.clear(); r = p.add_run("叶志伟（高级工程师）")
        r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
    elif i == 23 and '***' in t:
        p._element.getparent().remove(p._element)

# Fix dates - task book
for i, p in enumerate(doc.paragraphs):
    if i == 44:
        t = p.text.strip()
        if '日    期' in t or '期' in t:
            p.clear(); r = p.add_run("2024年5月31日")
            r.font.size=Pt(14); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)

print("Cover filled. Clearing template body...")

# ======== CLEAR TEMPLATE BODY ========
# Find body start (after KEY WORDS area)
body_start = -1
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "KEY WORDS" in t and len(t) < 4:
        body_start = i; break
if body_start == -1: body_start = 60

# Remove all paragraphs from body_start to end
total = len(doc.paragraphs)
for idx in range(total-1, body_start-1, -1):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
# Remove all tables
for t in list(doc.tables):
    t._element.getparent().remove(t._element)

print(f"Body cleared. Writing chapters...")

# Write Chinese + English abstracts into template placeholders
# Chinese abstract at paras 52-56 range, English at 58-66
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Chinese abstract body
    if i == 54 and len(t) > 200:
        p.clear(); r = p.add_run("基于Spring Cloud Alibaba微服务架构，设计并实现了面向郑州轻工业大学的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成果管理服务、文件服务和定时任务服务六个微服务组成，前端基于Vue 2和Element UI构建。论文首先分析了高校教学成果管理的业务需求，明确了教师、审核员和管理员三类角色的功能边界；随后从系统架构、功能模块、数据库、安全与权限四个方面阐述了总体设计方案；接着详细描述了网关认证过滤器、成果状态控制、审核流程事务保障等核心机制的实现方式；最后通过场景测试验证了系统的功能正确性。系统实现了教学成果从在线申报、审核流转到门户公示的全流程数字化管理。")
        rf(r, "宋体", sz=Pt(12))
    # Keywords
    if i == 56 and '关键' in t:
        p.clear(); r = p.add_run("关键词：教学成果管理；微服务架构；Spring Cloud；审核流程；Vue 2")
        rf(r, "宋体", sz=Pt(12))
    # English abstract body
    if i == 64 and len(t) > 200:
        p.clear(); r = p.add_run("Based on the Spring Cloud Alibaba microservice architecture, this thesis designs and implements a teaching achievement management system for Zhengzhou University of Light Industry. The system comprises six microservices: gateway, authentication center, system management, achievement management, file service, and scheduled task service, with the frontend built on Vue 2 and Element UI. The thesis first analyzes the business requirements of university teaching achievement management, clarifying the functional boundaries of teachers, auditors, and administrators. It then elaborates on the overall design including system architecture, functional modules, database, and security mechanisms. The implementation details of gateway authentication filter, achievement status control, and audit workflow transaction guarantee are subsequently presented. Functional scenario testing validates the correctness of core modules. The system achieves full-process digital management of teaching achievements from online submission, audit workflow, to portal publication.")
        rf(r, "宋体", sz=Pt(12))
    # KEY WORDS
    if i == 66 and 'KEY WORDS' in t:
        p.clear(); r = p.add_run("KEY WORDS: Teaching Achievement Management; Microservices Architecture; Spring Cloud; Audit Workflow; Vue 2")
        r.font.size=Pt(12); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)

# Remove template hints
for p in list(doc.paragraphs):
    t = p.text.strip()
    if ('小四' in t or '小三' in t) and '行距' in t and len(t) < 30:
        p._element.getparent().remove(p._element)

print("Abstracts written. Writing body...")

# ======== CHAPTER 1 ========
h1(doc, "第一章 引言")

h2(doc, "1.1 选题背景")
body(doc, "近年来，高等教育教学改革持续深化。教育部相继发布《关于深化本科教育教学改革全面提高人才培养质量的意见》等指导性文件，明确要求完善教学成果评价机制，推动教学管理工作的数字化转型 [12]。在此背景下，教学成果的申报、评审与统计已成为衡量教师教学水平、引导教改方向的重要工作环节。")
body(doc, "郑州轻工业大学现有二十余个教学单位，涵盖工学、理学、管理学、文学、艺术学等多个学科门类，每年定期组织校级教学成果奖的申报与评审工作，涉及教改项目、教材建设、课程建设、竞赛指导等多种类别。传统的管理方式以纸质材料提交与人工汇总为主，存在以下突出问题：成果信息分散于各学院教务办公室，缺乏统一的数据汇聚与查询入口；审核流转依赖纸质签字或邮件转发，过程不够透明且周期偏长；成果统计依赖手工制表，按学院、类型、等级等多维度进行汇总分析的效率较低；已通过审核的优秀成果缺少对外展示窗口，不利于教学经验的传播与示范引领。")
body(doc, "与此同时，软件架构技术的发展为教学管理信息化建设提供了新的思路。Spring Cloud Alibaba 微服务体系为构建高内聚、低耦合的分布式应用提供了成熟的技术组件：Nacos 实现服务注册发现与统一配置管理，Spring Cloud Gateway 作为 API 网关统一入口并集中执行认证与鉴权 [7]，Redis 支持令牌缓存与会话状态管理 [10]，MyBatis 提供灵活的数据库访问映射能力 [9]。采用微服务架构将系统拆分为认证中心、系统管理、成果管理、文件服务等独立模块，各服务可独立开发、部署与扩展，较好地适应教学管理业务持续演进的需求。")
body(doc, "基于上述背景，本课题设计并实现了面向郑州轻工业大学的教学成果管理系统。系统基于 Spring Cloud 微服务架构，支持教师在线申报教学成果、校级审核流转、多维度统计分析以及成果门户公示等功能，为高校教学成果的全流程数字化管理提供了一套实用的解决方案。")

h2(doc, "1.2 国内外研究现状")
body(doc, "教学成果管理的信息化研究可以追溯到教育管理信息系统（EMIS）的早期建设阶段。20世纪90年代，欧美高校率先在校园网基础上建立了面向教务和科研的综合管理平台。2014年，James Lewis和Martin Fowler系统阐述了微服务架构风格，提出将单一应用程序划分为一组小型服务，每个服务围绕特定业务能力构建，通过轻量级通信机制协作，独立部署和扩展 [1]。这一理念深刻影响了后续教育信息化系统的技术选型。")
body(doc, "进入21世纪后，随着Java EE和.NET等企业级开发框架的普及，具有完整业务流程的教学管理系统开始在教育领域大规模应用。Chris Richardson在其著作中详细总结了微服务架构的实现模式，包括服务拆分策略、跨服务通信、数据一致性处理以及分布式事务等关键问题的解决方案 [2]。Sam Newman则从工程实践角度出发，探讨了微服务系统的构建、测试、部署和运维方法 [3]。")
body(doc, "国内方面，高校教学管理信息化起步相对较晚但进展迅速。从2003年教育部启动精品课程建设工作开始，各高校陆续建设了不同形态的教务管理系统。早期以清华大学教育技术研究所开发的综合教务管理系统为典型代表，该类系统采用C/S架构，侧重于学籍管理、排课选课和成绩管理等教务流程，对教学成果的专项管理覆盖较为薄弱。2010年以后，以正方教务系统和青果教务系统为代表的B/S架构平台在众多高校推广使用，这些系统虽然增加了教研项目管理功能，但成果申报多以表单形式嵌入教务流程，缺乏独立的成果审核工作流和统计分析模块。")
body(doc, "近年来，微服务架构在教育领域的应用成为研究热点。黄建平在JWT与Redis结合的单点登录系统研究中，提出了一种基于双令牌的认证机制 [4]。孙卫琴在《精通Spring Cloud Alibaba》中系统介绍了微服务架构下的服务注册、配置管理和远程调用等核心技术 [5]。王志刚等人在Spring Cloud微服务实战中，结合具体案例展示了微服务架构在复杂业务场景中的应用方法 [6]。")
body(doc, "综合以上研究现状，现有教学管理系统在教学成果管理方面仍存在以下不足：第一，多数平台仅将成果管理作为教务系统的附属模块，未能提供独立的成果全生命周期管理功能，包括从草稿编辑、正式提交、审核流转到门户公示的完整闭环；第二，对多学院环境下数据权限隔离的关注不足，缺乏按部门层级精细化控制数据访问范围的能力 [20]；第三，部分已有研究在技术方案上仍停留在理论探讨或原型验证阶段，与工程化落地的实际系统之间存在较大差距。本课题在上述研究基础上，针对郑州轻工业大学多学院协同管理的实际需求，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。")

h2(doc, "1.3 研究内容")
body(doc, "本课题以郑州轻工业大学教学成果管理的实际需求为导向，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成就管理服务、文件服务和定时任务服务六个微服务组成，配合Vue 2前端框架实现用户交互。具体研究内容包括以下四个方面：")
body(doc, "（1）教学成果申报与管理功能。为教师提供在线申报教学成果的完整功能，支持富文本内容编辑、附件上传、草稿保存与正式提交。教师可在成果列表中查看自己所有申报记录，支持按状态、类型等条件筛选，并在成果被驳回后重新编辑并提交，或在审核过程中主动撤回申报。管理员具有所有成果的查看、修改和删除权限，并支持数据的批量导出。")
body(doc, "（2）多级审核流程。系统实现了教学成果的状态机流转逻辑，包含草稿（0）、审核中（2）、已通过（3）和已驳回（4）四种状态。教师提交申报后成果进入审核状态，审核员可对待审成果进行通过或驳回处理，驳回时需填写审核意见以便教师了解原因。被驳回的成果允许教师修改后重新提交，审核中的成果允许教师主动撤回，已通过的成果则不允许再次修改，以保证审核结果的权威性。")
body(doc, "（3）成果统计与分析。系统提供多维度的统计分析能力，支持按成果状态、成果类型、所属学院等维度进行汇总统计，同时提供总数、今日新增、本周新增、本月通过数和本月驳回数等实时指标。教师端也提供个人成果的状态分布统计，便于教师了解自己的申报进度。")
body(doc, "（4）门户信息公示。系统内置新闻管理模块，支持发布教学动态和通知公告，并提供对外的门户展示页面。审核通过的教学成果可自动推送到门户网站进行公示，支持按浏览量和发布时间排序，便于师生查看和参考优秀成果。门户内容为公开访问，无需登录即可查看。")

h2(doc, "1.4 论文结构安排")
body(doc, "本文共分为六章，各章内容安排如下：")
body(doc, "第一章为引言。介绍课题的选题背景与意义，回顾国内外教学管理信息化和微服务架构在教育领域应用的研究现状，明确本课题的研究内容和论文组织结构。")
body(doc, "第二章为需求分析。在对系统进行概述的基础上，分析系统涉及的用户角色（超级管理员、审核员、教师），从功能和非功能两个维度详细描述系统需求，功能需求涵盖成果管理、审核流程、门户新闻、统计分析和用户权限管理五个核心模块，非功能需求涵盖安全性、可靠性和可维护性等方面。")
body(doc, "第三章为总体设计。从系统架构、功能模块划分、数据库、安全与权限以及技术栈选型五个层面阐述系统的整体设计方案。其中数据库设计部分包含E-R图和核心表的字段说明，安全设计部分详细说明JWT+Redis双令牌认证机制和RBAC权限控制方案。")
body(doc, "第四章为详细实现。逐一阐述网关与认证模块、成果管理模块、审核流程模块、门户展示模块、文件服务和前端实现的具体技术细节，包括核心代码逻辑、接口设计、关键配置和业务流程的实现方式。")
body(doc, "第五章为测试。介绍系统的测试环境和测试方法，从功能验证角度对成果申报、审核流转、统计分析等核心功能进行场景测试，并给出测试结果。")
body(doc, "第六章为总结与展望。总结本课题的工作内容和实际完成情况，指出系统当前存在的不足，并对后续改进方向进行展望。")

print("Chapter 1 done. Continuing...")
doc.save(OUTPUT)

h1(doc, "第二章 需求分析")
h2(doc, "2.1 系统概述")
body(doc, "高校教学成果管理系统是一个面向郑州轻工业大学各教学单位的网络化服务平台，服务对象包括教师、学院审核员、校级审核员和系统管理员。系统的核心业务是对教师的教学成果申报进行全流程数字化管理，从教师填写成果内容并上传证明材料开始，经过审核员的通过或驳回处理，最终形成可查询、可统计、可展示的数字化成果档案。")
body(doc, "从业务角度看，系统涵盖了教学成果管理的全生命周期。教师在成果管理模块中可以新建成果、保存草稿、编辑内容、上传附件，也可以将完成编辑的成果提交审核。审核员通过审核模块查看待审成果的详情，填写审核意见并做出通过或驳回的决定。审核通过的成果会自动推送至门户网站进行公示。管理员可以查看全校各学院的成果申报情况，利用统计分析功能掌握整体进度。")

h2(doc, "2.2 角色分析")
body(doc, "根据教学成果管理的实际业务流程，系统设置以下三类用户角色，各角色的权责和操作范围明确区分：")
body(doc, "（一）普通教师。教师是系统的主要使用者，负责编写和提交自己的教学成果申报。教师可以新建成果记录、保存草稿、修改内容、上传证明材料附件，并在确认无误后将成果提交审核。对于处于审核中状态的成果，教师可执行撤回操作将其退回草稿状态；对于被驳回的成果，教师可根据审核意见修改后重新提交。已通过审核的成果不允许再次编辑，以保证审核结果的不可篡改性。教师的所有操作被限定在本人的成果范围内，系统通过强制绑定登录用户ID的方式确保数据隔离。")
body(doc, "（二）审核员。审核员拥有查看待审成果、执行审核操作和查看审核历史记录的权限。审核员在审核界面可以查看成果的完整内容、附件材料和教师信息，然后填写审核意见并选择通过或驳回。审核过程全程留痕，每一次审核操作都会生成一条审核记录，包含审核人、审核时间、审核结果和审核意见。系统当前采用单级校级审核模式，但在数据库设计中预留了审核级别字段（audit_level），可以在后续版本中扩展为院级初审和校级终审的两级审核模式。")
body(doc, "（三）超级管理员。超级管理员拥有系统的最高权限，可以查看和管理全校所有教师的成果数据，不受数据权限的限制。管理员还可以管理系统用户、角色、菜单、部门、字典等基础数据，配置权限分配和数据访问范围策略，以及查看操作日志和登录记录。系统设计中，超级管理员的用户ID为固定值1，在数据权限切面中绕过所有数据范围过滤，确保其可以全局视图管理数据。")
body(doc, "三类角色的权限体系通过基于RBAC（角色基于访问控制）的权限管理模块来实现。每个菜单项和接口都绑定了特定的权限标识符，用户通过角色获得对应的权限集合。在后端接口层面，通过@RequiresPermissions注解和切面编程在请求进入业务逻辑前执行权限校验，确保未授权的操作无法执行。")

print("Chapter 2 done. Continuing...")
doc.save(OUTPUT)

h2(doc, "2.3 功能需求")
h3(doc, "2.3.1 成果管理")
body(doc, "成果管理模块提供教学成果的增删改查基础操作，是系统最核心的功能区域。教师通过成果管理模块可以新建成果记录，填写成果标题、内容详述、选择成果类型和申报等级，并上传证明材料附件。成果内容支持富文本编辑，教师可以在成果详述中插入图片、表格等富文本元素。")
body(doc, "系统对成果的编辑操作实施了严格的状态控制。处于审核中状态（2）或已通过状态（3）的成果不允许教师修改，以保证审核过程的严肃性和审核结果的不可篡改性。草稿状态（0）和已驳回状态（4）的成果允许教师自由编辑。教师只能查看和操作自己的成果，系统在控制器层通过SecurityUtils.getUserId()强制绑定当前登录用户ID，确保每个教师只能访问自己的数据。")
body(doc, "管理员端提供全校成果的统一查看界面，支持按成果标题、教师姓名、学院、状态、类型等条件进行复合查询。管理员可以修改任意成果的内容和状态，也可以执行删除操作。系统还提供了数据导出功能，支持将成果列表导出为Excel文件，便于数据备份和离线分析。")

h3(doc, "2.3.2 审核流程")
body(doc, "审核流程模块负责教学成果的状态流转和审核记录管理。成果的状态设计为四种：草稿（0）、审核中（2）、已通过（3）和已驳回（4）。状态之间的转换遵循固定的规则：教师新建成果时默认为草稿状态，也可以直接提交为审核中状态；处于审核中状态的成果经审核员处理后可转为已通过或已驳回；被驳回的成果允许教师修改后重新提交审核；审核中的成果允许教师主动撤回为草稿。已通过的成果不允许任何修改，以保证审核结果的终局性。")
body(doc, "审核员在审核界面可以查看待审成果的完整信息，包括成果标题、内容详述、附件材料、教师信息和所属学院。审核员需要填写审核意见，然后选择通过或驳回。每一次审核操作都会在数据库中保存一条审核记录，包含审核人ID、审核人姓名、审核级别、审核结果、审核意见和审核时间。审核记录与成果通过achievement_id关联，支持查询某个成果的历史审核轨迹。")
body(doc, "系统当前采用单级审核模式，审核流程由教师提交后直接进入校级审核环节。数据库设计中预留了audit_level字段用于标识审核级别，这为后续扩展为院级初审加校级终审的两级审核模式提供了数据结构支撑。")

h3(doc, "2.3.3 门户新闻")
body(doc, "门户新闻模块提供教学动态和通知公告的发布与展示功能。管理员可以发布新闻动态和通知公告两类内容，支持富文本正文编辑、封面上传、摘要填写和排序权重设置。新闻发布后可以设置是否推送到门户网站进行公开展示。")
body(doc, "门户前端提供对外的新闻列表和新闻详情页面，支持按发布时间和浏览量排序。门户内容为公开访问，无需登录即可查看。这一设计通过网关层的白名单机制实现：网关在AuthFilter中检测URL前缀，当请求路径以/achievement/portal/news开头时，直接放行而不执行令牌验证，从而实现公共内容的无障碍访问。")

h3(doc, "2.3.4 统计分析")
body(doc, "统计分析模块提供多维度的成果数据汇总能力，帮助管理员掌握全校教学成果申报的整体情况。统计维度包括：按成果状态统计，可以查看草稿、审核中、已通过和已驳回各状态的数量分布；按成果类型统计，分别汇总科研指导、教材建设、竞赛指导、教学改革、教学评估和优秀课程六种类型的数据；按学院统计，展示各学院成果申报数量的排名情况。统计数据通过MyBatis映射执行SQL聚合查询获取，服务层提供了countByStatus、countByCategory、countByCollege等聚合方法。")
body(doc, "除了基础的分类统计，系统还提供了多个关键指标的实时查询，包括成果总数、待审核数量、今日新增、本周新增、本月通过数和本月驳回数。这些指标可以在管理员首页以数据卡片的形式直观展示，便于快速了解当前成果申报的动态情况。教师端同样提供个人维度的状态统计，便于教师了解自己的成果申报进展。")

h3(doc, "2.3.5 用户权限管理")
body(doc, "用户权限管理模块提供基于RBAC模型的权限控制功能，包括用户管理、角色管理、菜单管理和部门管理四个子模块。用户管理支持用户的新增、编辑、删除、密码重置和状态控制；角色管理支持角色的创建、权限分配和用户关联；菜单管理支持菜单树的动态维护和权限标识符绑定；部门管理支持学校-学院-系部的多级组织结构维护。")
body(doc, "系统在权限控制方面实现了两个层级的检查：接口级权限校验和数据级权限过滤。接口级权限通过@RequiresPermissions注解和面向切面编程实现，在控制器方法执行前检查用户是否拥有对应的权限字符串。数据级权限通过@DataScope注解实现，该注解会在MyBatis查询时动态注入数据范围条件，确保不同学院的教师和审核员只能看到权限范围内的数据。数据范围分为五个级别：全部数据、自定义数据、本部门数据、本部门及以下数据和仅本人数据，可以根据实际需要为不同角色配置不同的数据访问范围。")

h2(doc, "2.4 非功能需求")
body(doc, "非功能需求定义了系统在可用性、安全性、可维护性和兼容性等方面应达到的质量属性。结合高校教学成果管理的业务特点，本系统的非功能需求主要包括以下几个方面：")
body(doc, "安全性方面。系统需要对所有接口请求进行身份验证，确保未登录用户无法访问受保护的业务接口。认证采用JWT与Redis结合的双令牌机制：JWT承载用户基本身份信息（用户ID、用户名、用户密钥），Redis存储完整的登录状态和权限信息，既保证了令牌的自包含性，又实现了服务端的会话控制。网关层需要过滤非法请求，防止越权访问。微服务之间的内部调用需要通过内部认证机制来保证调用合法性，避免外部请求伪装成内部调用。")
body(doc, "可靠性方面。系统需要保证核心业务操作的数据一致性，特别是成果状态转换和审核操作需要具有事务保障。当审核员对成果执行审核时，更新成果状态和插入审核记录应保持原子性，要么同时成功，要么同时失败，不能出现状态已更新但审核记录未生成的不一致情况。网关层需要对异常请求进行统一拦截和响应，避免将服务端错误直接暴露给前端。")
body(doc, "可维护性方面。各微服务需要保持独立性，单个服务的修改和重新部署不应影响其他服务的正常运行。配置管理采用Nacos统一配置中心，支持配置的动态刷新，避免因配置修改而重启服务。操作日志需要完整记录用户的关键操作，便于问题追踪和安全审计。")
body(doc, "兼容性方面。系统需要支持主流浏览器（Chrome、Firefox、Edge）的访问，前端页面需要在不同分辨率下正常显示。接口需要遵循RESTful设计风格，便于前后端分离开发和后续功能扩展。")

print("Ch2 done. Writing Ch3...")
doc.save(OUTPUT)

# ======== CHAPTER 3 ========
h1(doc, "第三章 总体设计")

h2(doc, "3.1 系统架构设计")
body(doc, "本系统采用基于Spring Cloud Alibaba的微服务架构，将业务功能拆分为多个独立的微服务，每个服务围绕单一业务领域构建，服务之间通过HTTP协议和Feign客户端进行通信。系统整体架构分为四个层次：用户层、网关层、业务服务层和数据层。")
body(doc, "用户层基于Vue 2框架构建，采用Element UI组件库实现前端页面。前端通过/dev-api前缀代理到网关服务（端口8080），网关自动剥离该前缀后转发到对应的后端服务。前端开发服务器配置了代理规则，将所有/dev-api请求转发至localhost:8080。")
body(doc, "网关层由Spring Cloud Gateway实现，作为系统的唯一入口，承担请求路由、身份验证、越权拦截和跨域处理等职责。网关通过AuthFilter全局过滤器对每个请求进行令牌验证，合法请求会被添加用户身份Header后转发到下游的业务服务。网关还配置了跨域策略，允许前端页面通过浏览器直接调用后端接口。")
body(doc, "业务服务层包含认证中心（edu-auth，端口9200）、系统管理服务（edu-system，端口9201）、成果管理服务（edu-achievement，端口9205）、文件服务（edu-file，端口9300）和定时任务服务（edu-job，端口9203）五个微服务。各服务通过Nacos进行服务注册与发现，配置信息统一存储在Nacos配置中心。服务启动顺序有严格要求：认证中心必须在网关之前启动，因为网关在启动时就会通过Redis验证令牌有效性，如果认证服务未启动，登录流程将无法正常工作。")
body(doc, "数据层包括MySQL数据库、Redis缓存和文件对象存储。MySQL负责持久化存储业务数据和用户信息，Redis负责登录会话缓存和在线用户管理，文件服务支持本地存储、MinIO和FastDFS三种后端，可通过配置切换。系统还配置了监控中心（edu-monitor，端口9100），基于Spring Boot Admin实现各微服务的健康检查和运行状态监控。")

h2(doc, "3.2 功能模块划分")
body(doc, "根据需求分析，系统功能模块如图3-2所示，主要包括以下六个模块：")
body(doc, "（1）认证与权限模块。包括用户登录、登出、令牌刷新、用户注册等功能，由edu-auth服务和网关AuthFilter共同实现。登录时用户提交用户名和密码，认证服务验证凭证后生成JWT令牌，并将完整的LoginUser对象存储在Redis中，令牌过期时间为720分钟。")
body(doc, "（2）系统管理模块。包括用户管理、角色管理、菜单管理、部门管理、字典管理、通知管理和操作日志等功能，由edu-system服务承担，提供整个系统运行所需的基础数据支持。")
body(doc, "（3）成果管理模块。包括教师成果申报、成果列表查询、成果编辑、附件上传、Excel导出等功能，由edu-achievement服务承担，是系统的核心业务模块。")
body(doc, "（4）审核流程模块。包括待审成果查询、审核通过/驳回、审核意见填写、审核记录查询等功能，由edu-achievement服务中的EduAuditRecordController承担。审核操作会同时更新成果状态和插入审核记录。")
body(doc, "（5）门户展示模块。包括新闻动态和通知公告的发布/编辑和门户展示，由edu-achievement服务中的EduNewsController和PortalNewsController承担，门户内容为公开访问，无需登录。")
body(doc, "（6）文件服务模块。提供附件文件的上传、下载和前端静态资源的访问功能，由edu-file服务承担，支持本地存储、MinIO和FastDFS三种后端，可通过配置切换。")

h2(doc, "3.3 数据库设计")
h3(doc, "3.3.1 E-R模型")
body(doc, "系统的数据库设计围绕教学成果管理的核心业务展开，主要实体包括教师（sys_user）、学院（sys_dept）、教学成果（edu_achievement）、审核记录（edu_audit_record）和门户新闻（edu_news）。实体之间的关系如下：一个教师可以申报多个教学成果，一个学院下有多个教师，一个教学成果可以有多条审核记录，审核通过的成果可以发布到门户展示。系统E-R图如图3-3所示。")

h3(doc, "3.3.2 核心表结构")
body(doc, "系统核心业务表包括教学成果主表（edu_achievement）、审核记录表（edu_audit_record）和门户新闻表（edu_news），其设计结构如下各表所示。系统管理相关表包括用户表（sys_user）、角色表（sys_role）、菜单表（sys_menu）、部门表（sys_dept）、字典数据表（sys_dict_data）、字典类型表（sys_dict_type）、参数配置表（sys_config）和操作日志表（sys_oper_log）。")

table_cap(doc, "表3-1 教学成果主表(edu_achievement)")
add_tbl(doc, ["字段名","类型","说明"],
    [["achievement_id","bigint","主键ID，自增"],["title","varchar(200)","成果标题"],["content","text","成果内容(富文本HTML)"],
     ["file_url","varchar(1500)","证明材料文件路径"],["teacher_id","bigint","教师ID，关联sys_user"],
     ["college_id","bigint","学院ID，关联sys_dept"],["status","char(1)","0草稿/2审核中/3已通过/4已驳回"],
     ["category","char(1)","成果类型(1~6)"],["level","varchar(30)","申报等级"],
     ["create_time","datetime","创建时间"],["del_flag","char(1)","删除标志(0存在2删除)"]])

table_cap(doc, "表3-2 审核记录表(edu_audit_record)")
add_tbl(doc, ["字段名","类型","说明"],
    [["record_id","bigint","主键ID，自增"],["achievement_id","bigint","关联成果ID"],
     ["audit_level","char(1)","审核级别"],["audit_result","char(1)","1通过/2驳回"],
     ["audit_opinion","varchar(500)","审核意见"],["auditor_id","bigint","审核人ID"],
     ["auditor_name","varchar(64)","审核人姓名"],["create_time","datetime","审核时间"]])

table_cap(doc, "表3-3 门户新闻表(edu_news)")
add_tbl(doc, ["字段名","类型","说明"],
    [["news_id","bigint","主键ID，自增"],["title","varchar(120)","新闻标题"],
     ["summary","varchar(255)","新闻摘要"],["cover_image","varchar(500)","封面图URL"],
     ["content","longtext","富文本正文"],["publish_time","datetime","发布时间"],
     ["view_count","bigint","浏览量"],["publish_portal","char(1)","是否发布到门户"],
     ["sort_weight","int","排序权重"],["notice_type","char(1)","1通知公告/2新闻动态"]])

table_cap(doc, "表3-4 系统用户表(sys_user)")
add_tbl(doc, ["字段名","类型","说明"],
    [["user_id","bigint","主键ID，自增"],["user_name","varchar(30)","登录账号"],
     ["nick_name","varchar(30)","用户昵称"],["password","varchar(100)","密码(BCrypt加密)"],
     ["phonenumber","varchar(11)","手机号码"],["email","varchar(50)","电子邮箱"],
     ["sex","char(1)","性别(0男1女2未知)"],["dept_id","bigint","部门ID"],
     ["status","char(1)","状态(0正常1停用)"],["del_flag","char(1)","删除标志"],
     ["create_time","datetime","创建时间"]])

table_cap(doc, "表3-5 系统角色表(sys_role)")
add_tbl(doc, ["字段名","类型","说明"],
    [["role_id","bigint","主键ID，自增"],["role_name","varchar(30)","角色名称"],
     ["role_key","varchar(100)","角色权限字符串"],["role_sort","int","显示顺序"],
     ["data_scope","char(1)","数据范围(1~5)"],["status","char(1)","角色状态(0正常)"],
     ["create_time","datetime","创建时间"]])

table_cap(doc, "表3-6 系统菜单表(sys_menu)")
add_tbl(doc, ["字段名","类型","说明"],
    [["menu_id","bigint","主键ID，自增"],["menu_name","varchar(50)","菜单名称"],
     ["parent_id","bigint","父菜单ID"],["perms","varchar(100)","权限标识符"],
     ["menu_type","char(1)","类型(M目录C菜单F按钮)"],["path","varchar(200)","路由地址"],
     ["component","varchar(255)","前端组件路径"],["icon","varchar(100)","菜单图标"],
     ["order_num","int","显示顺序"],["status","char(1)","菜单状态(0显示1隐藏)"]])

table_cap(doc, "表3-7 字典数据表(sys_dict_data)")
add_tbl(doc, ["字段名","类型","说明"],
    [["dict_code","bigint","主键ID"],["dict_sort","int","排序"],
     ["dict_label","varchar(100)","字典标签"],["dict_value","varchar(100)","字典键值"],
     ["dict_type","varchar(100)","字典类型"],["status","char(1)","状态"]])

table_cap(doc, "表3-8 字典类型表(sys_dict_type)")
add_tbl(doc, ["字段名","类型","说明"],
    [["dict_id","bigint","主键ID"],["dict_name","varchar(100)","字典名称"],
     ["dict_type","varchar(100)","字典类型(unique)"],["status","char(1)","状态"]])

table_cap(doc, "表3-9 参数配置表(sys_config)")
add_tbl(doc, ["字段名","类型","说明"],
    [["config_id","int","主键ID"],["config_name","varchar(100)","参数名称"],
     ["config_key","varchar(100)","参数键名"],["config_value","varchar(500)","参数键值"],
     ["config_type","char(1)","系统内置(Y/N)"]])

table_cap(doc, "表3-10 操作日志表(sys_oper_log)")
add_tbl(doc, ["字段名","类型","说明"],
    [["oper_id","bigint","主键ID"],["title","varchar(50)","模块标题"],
     ["business_type","int","业务类型"],["method","varchar(200)","方法名称"],
     ["oper_name","varchar(50)","操作人员"],["oper_time","datetime","操作时间"]])

print("Ch3 tables done. Continuing design sections...")
doc.save(OUTPUT)

h2(doc, "3.4 安全与权限设计")
body(doc, "系统的安全设计采用了多层防护体系，从网关层、服务层和数据层三个层面确保系统安全。")
body(doc, "网关层安全通过AuthFilter全局过滤器实现，该过滤器的执行顺序为-200，优先于其他业务过滤器。对于每个请求，首先检查URL是否在白名单中（包括登录接口、验证码接口和门户新闻接口），白名单中的请求直接放行；其余请求需要携带有效的JWT令牌。过滤器从HTTP请求头中提取Authorization字段，解析JWT获取用户密钥，然后在Redis中检查login_tokens:{user_key}是否存在，如果不存在则认为登录状态已过期。验证通过后，过滤器会将user_key、user_id和username三个Header添加到请求中，并移除可能存在的from-source Header，防止外部请求伪装成内部调用。")
body(doc, "服务层安全包括接口权限和数据权限两个维度。接口权限通过@RequiresPermissions、@RequiresRoles和@RequiresLogin注解实现，由PreAuthorizeAspect切面在方法执行前进行权限校验。数据权限通过@DataScope注解实现，在MyBatis查询时动态注入数据范围条件。数据权限分为五个级别：级别1为全部数据权限，级别2为自定义数据权限，级别3为本部门数据权限，级别4为本部门及以下数据权限，级别5为仅本人数据权限。超级管理员（userId=1）在数据权限切面中被特殊处理，绕过所有数据范围过滤。")
body(doc, "微服务间的内部调用通过Feign客户端实现，FeignRequestInterceptor会自动向请求中添加Authorization、user_id、user_key和username等Header。接口层通过@InnerAuth注解标记内部接口，InnerAuthAspect切面会校验from-source Header是否为inner，确保只有内部服务可以调用这些接口。")

h2(doc, "3.5 技术栈选型")
body(doc, "本系统技术栈选型以成熟稳定和与微服务架构的配套程度为主要考量。后端基于Java 1.8和Spring Boot 2.7.18构建，使用Spring Cloud 2021.0.9和Spring Cloud Alibaba 2021.0.6.1作为微服务基础设施。Nacos 2.2.3提供服务注册发现和配置管理能力。数据持久化层采用MySQL 8.0作为主数据库，MyBatis作为ORM框架，PageHelper实现分页查询，Druid作为数据库连接池。缓存层采用Redis 5+，存储登录会话和业务缓存数据。")
body(doc, "前端基于Vue 2框架和Element UI组件库构建单页面应用，采用基于RuoYi-Vue的二次开发模式，在其基础上扩展了教学成果管理相关的页面和组件。前端通过RESTful API与后端通信，数据格式采用JSON。前端开发服务器配置了代理规则，实现前后端分离开发。")
body(doc, "其他关键组件包括：Alibaba TransmittableThreadLocal解决异步任务中上下文传递问题，确保@Async方法中仍可获取当前登录用户信息；Quartz实现定时任务调度；SpringDoc OpenAPI自动生成接口文档；Apache POI处理Excel数据导出；MinIO SDK实现对象存储文件操作。整体技术栈完全基于开源组件，没有商业许可依赖。")

print("Chapter 3 done. Writing Ch4...")
doc.save(OUTPUT)

h1(doc, "第四章 详细实现")

h2(doc, "4.1 开发环境")
body(doc, "系统开发环境基于Windows 11平台搭建，主要工具链如下：后端使用JDK 1.8和Maven 3.6+进行项目构建与依赖管理，开发IDE采用IntelliJ IDEA；前端使用Node.js 14+和npm/pnpm作为包管理器，开发服务器端口为80；数据库采用MySQL 8.0.33，缓存采用Redis 5.0，服务注册与配置中心采用Nacos 2.2.3单机模式。")
body(doc, "项目内置两套数据库初始化脚本：edu_system.sql包含全部业务表结构及示例数据，edu_config.sql为Nacos配置数据库的初始化脚本。服务启动顺序必须严格遵守，这是本系统区别于标准微服务项目的关键约束：Nacos和Redis启动后，依次启动edu-auth、edu-gateway、edu-system、edu-achievement、edu-file。其中edu-auth必须在edu-gateway之前启动，因为网关在启动时便会通过Redis进行令牌校验，如果认证服务未启动，登录流程将无法正常工作。前端执行npm run dev即可启动开发服务器，页面通过http://localhost访问。编译打包使用mvn clean package -DskipTests命令。")

h2(doc, "4.2 网关与认证模块")
body(doc, "网关模块基于Spring Cloud Gateway构建，承担统一入口、请求路由、身份验证和跨域处理四项核心职责。网关配置了对各下游服务的路由规则，例如/auth/**路径转发至edu-auth服务，/system/**路径转发至edu-system服务，/achievement/**路径转发至edu-achievement服务。所有路由规则的服务地址均采用Nacos服务名(lb://edu-auth等)，实现服务名到实际地址的动态解析。")
body(doc, "身份验证通过AuthFilter全局过滤器实现，是网关模块的核心组件。AuthFilter实现了GlobalFilter和Ordered接口，执行顺序设置为-200，优先于其他业务过滤器。其核心逻辑如下：首先检查请求URL是否以/achievement/portal/news开头，若是则直接放行——这是系统中硬编码的公共访问通道，为门户展示提供无障碍服务。其次，通过IgnoreWhiteProperties配置类检查URL是否在白名单中，白名单包含登录接口、验证码接口、注册接口等无需认证的接口。对于非白名单请求，从HTTP请求头中提取Authorization字段，移除Bearer前缀后解析JWT令牌，获取user_key、user_id和username。随后在Redis中检查login_tokens:{user_key}是否存在，若不存在则返回401未授权响应。最后将三个用户信息Header添加到转发请求中，并删除from-source Header，确保下游服务不会误判为内部调用。")
body(doc, "认证服务（edu-auth）提供登录、登出、令牌刷新和用户注册四个接口。登录接口/token/login接收LoginBody参数，包含username和password，SysLoginService负责验证用户凭证。验证通过后，生成一个UUID作为会话标识（user_key），构建LoginUser对象包含用户信息、角色列表和权限集合，存储到Redis的login_tokens:{user_key}键下，设置TTL为720分钟。本系统采用双令牌设计：JWT令牌仅承载user_key、user_id和username三个轻量字段，不包含权限信息；完整的用户权限信息存储在Redis中。这种设计使得即使JWT令牌被截获，攻击者也无法获取用户的完整权限集合，而服务端可以通过删除Redis中的会话数据来主动强制用户下线。JWT的签名密钥为固定字符串，解析时通过JwtUtils类完成。")
body(doc, "下游服务在处理请求时，通过HeaderInterceptor拦截器从请求头中提取user_id、user_key和username，封装为LoginUser对象并设置到SecurityContextHolder中。业务代码可以随时通过SecurityUtils.getUserId()、SecurityUtils.getUsername()等静态方法获取当前用户信息。SecurityContextHolder基于Alibaba的TransmittableThreadLocal实现，即使在@Async异步方法中也能正确获取当前用户上下文，这解决了传统ThreadLocal在异步任务中上下文丢失的问题。FeignRequestInterceptor会在Feign调用时自动向请求中添加Authorization、user_id、user_key和username等Header，确保内部服务调用链路中用户身份的完整传递。")

h2(doc, "4.3 成果管理模块")
body(doc, "成果管理模块是系统的核心业务模块，由EduAchievementController、EduAchievementServiceImpl和EduAchievementMapper三层组成。Controller层定义了面向管理员和教师两套接口。管理员接口包括/achievement/list（列表查询，支持按标题、教师姓名、学院、状态、类型等条件复合查询）、/achievement/export（Excel导出，支持按时间段过滤）、/achievement/{id}（单个查询）、POST新增、PUT修改和DELETE批量删除。所有接口均通过@RequiresPermissions注解绑定权限标识符，如achievement:achievement:list、achievement:achievement:edit等，确保不同角色只能访问其被授权的接口。")
body(doc, "管理员修改成果时，系统会先从数据库查出旧数据，然后判断当前用户是否为超级管理员。非超级管理员需要校验是否为成果所有者，且已通过状态（3）不允许修改。超级管理员可以修改任意成果的内容和状态，如果未传入新状态则保持原状态不变。这一设计兼顾了数据保护和管理灵活性。")
body(doc, "教师接口包括teacherAddAchievement（新增）、teacherListAchievement（列表查询）、teacherUpdateAchievement（修改）、teacherDelAchievement（删除）、teacherGetAchievement（单个查询）、teacherListAllAchievement（全量查询含状态统计）、teacherResubmit（驳回后重新提交）和teacherWithdraw（撤回审核）。所有教师接口在处理时都会调用SecurityUtils.getUserId()获取当前登录用户ID，并强制设置到查询条件中，确保教师只能看到和操作自己的成果数据。teacherListAllAchievement接口除返回列表数据外，还在内存中统计各状态数量（已通过、审核中、已驳回），并封装在AjaxResult中一并返回。")
body(doc, "教师修改成果时的状态控制逻辑如下：首先从数据库查出旧数据，校验成果是否存在、是否为本人所有。然后检查当前状态：审核中（2）或已通过（3）状态不允许教师修改；草稿（0）状态允许保存为草稿或提交为审核中；已驳回（4）状态允许修改内容但保持驳回状态。教师可以通过teacherResubmit接口将驳回的成果重新提交为审核中状态，提交时需要校验成果必须属于当前教师且状态为已驳回；也可以通过teacherWithdraw接口将审核中的成果撤回为草稿，撤回时仅允许状态为审核中的成果。")
body(doc, "服务层EduAchievementServiceImpl负责业务逻辑处理和数据库操作协调。新增成果时，如果未明确指定状态或状态值非法，自动设置为草稿（0）。修改成果时使用MyBatis动态SQL，只更新非空字段，避免了全字段覆盖带来的数据丢失风险。删除成果采用逻辑删除，将del_flag设置为2而非物理删除记录，保留了数据可追溯性。服务层还提供了多个统计聚合方法——countByStatus、countByCategory、countByCollege、countTotal、countTotalPending、countTodayNew、countWeekNew、countMonthPassed和countMonthRejected等，满足管理员仪表盘和统计页面的数据需求。对于教师端，服务层还提供了countByStatusForTeacher、countByCategoryForTeacher和countTotalForTeacher等个人维度的统计接口。")

print("Ch4 core sections done. Continuing expansion...")
doc.save(OUTPUT)

h2(doc, "4.4 审核流程模块")
body(doc, "审核流程模块由EduAuditRecordController和IEduAuditRecordService实现，负责审核记录的查询和审核操作。审核员通过/audit/record/list接口查询待审成果列表，通过/audit/record/approve接口执行审核通过，通过/audit/record/reject接口执行审核驳回。审核通过时，系统同时执行两个操作：将成果的status字段更新为3（已通过），并向edu_audit_record表插入一条审核记录，记录审核人ID、姓名、审核级别、审核结果和审核意见。审核驳回时类似，将status更新为4（已驳回）并记录审核意见。")
body(doc, "为确保数据一致性，状态更新和审核记录插入应在同一数据库事务中执行，要么同时成功，要么同时失败。审核记录表通过achievement_id与成果主表关联，支持查询某个成果的全部审核历史。审核记录包含丰富的字段信息：record_id为自增主键，achievement_id关联成果，audit_level标识审核级别（预留字段，当前统一为校级审核），audit_result标识通过（1）或驳回（2），audit_opinion存储审核意见文字，auditor_id和auditor_name记录审核人信息，create_time记录审核时间。表中对achievement_id和auditor_id分别建立了索引，优化常见查询场景的性能。系统当前采用单级审核模式，数据库设计中预留的audit_level字段为后续扩展为院级初审加校级终审的两级模式提供了数据结构支撑。")

h2(doc, "4.5 门户展示模块")
body(doc, "门户展示模块包含EduNewsController和PortalNewsController两个控制器。EduNewsController提供新闻的增删改查接口，支持管理员在后台管理新闻内容，包括新闻标题、摘要、封面图、正文（富文本）、发布时间、排序权重、类型（通知公告或新闻动态）和是否发布到门户等字段。PortalNewsController提供面向公众的新闻查询接口，支持列表分页查询和单个新闻详情查看。")
body(doc, "门户接口的路径为/achievement/portal/news，在网关的AuthFilter中，所有以此前缀开头的请求都直接放行，不执行令牌验证。这意味着任何用户无需登录即可浏览门户新闻和成果公示内容。前端通过Axios发送请求时不携带Authorization头，网关将其视为公共请求直接放行。门户前端页面支持按发布时间倒序排列新闻列表，含封面图、标题、摘要和发布时间等信息，点击标题进入详情页展示富文本正文内容，同时浏览量自动增加。后端查询接口按publish_portal=1、status=0和del_flag=0三个条件过滤，确保仅已发布且未删除的内容在门户可见。")

h2(doc, "4.6 文件服务")
body(doc, "文件服务（edu-file）提供附件文件的上传、下载和前端静态资源的访问功能。服务支持三种存储后端：本地文件系统、MinIO对象存储和FastDFS分布式文件系统，可通过修改Nacos中的edu-file-dev.yml配置文件实现存储后端切换，无需修改代码。默认配置为本地存储模式，文件保存在服务器指定目录下，通过HTTP直接提供下载链接。服务的前缀路径为/static，用于统一管理静态资源的访问入口。")
body(doc, "文件上传接口支持单文件上传和多文件上传，上传完成后返回文件的访问URL。教师在成果申报时可以将返回的URL存入成果的file_url字段，系统支持存储多个附件文件路径，以逗号分隔。服务还配置了防盗链机制，可通过referer表白名单限制文件的访问来源，防止文件被外部网站直接链接盗用。当前版本中防盗链功能默认关闭，可根据实际部署环境开启。文件服务也支持配置CORS跨域策略，允许前端通过浏览器直接上传文件到文件服务器而不经过网关中转。")

h2(doc, "4.7 安全与权限详细实现")
body(doc, "系统的安全体系由多个切面和拦截器协同实现。接口权限校验通过PreAuthorizeAspect切面实现，在控制器方法执行前，读取@RequiresPermissions注解中声明的权限字符串，与当前登录用户的权限集合进行匹配。若用户不具备所需权限，直接抛出异常并返回403响应。此外系统还支持@RequiresRoles注解用于角色级别校验，以及@RequiresLogin注解用于仅要求登录状态的接口。内部接口通过InnerAuthAspect切面保护，校验请求头中的from-source字段是否为inner，确保只有Feign调用可以访问标注了@InnerAuth的内部接口。")
body(doc, "数据权限通过DataScopeAspect切面和@DataScope注解实现。在服务方法上添加@DataScope注解后，切面会根据当前用户角色的data_scope属性构建SQL过滤条件，并将其注入到方法参数params.dataScope中。MyBatis映射XML中通过${params.dataScope}引用该条件，在SQL执行时动态追加数据范围限制。使用${}而非#{}是因为此处需要注入的是原始SQL片段而非参数化占位符。五级数据范围定义：级别1为全部数据权限，级别2为自定义数据权限，级别3为本部门数据权限，级别4为本部门及以下数据权限，级别5为仅本人数据权限。超级管理员（userId=1）在切面中被特殊处理，直接跳过所有数据范围过滤，确保其可以全局视图管理数据。")
body(doc, "操作日志记录通过LogAspect切面实现，对添加了@Log注解的接口方法自动记录操作日志。切面会记录调用的模块名称、业务类型（新增/修改/删除/导出等）、调用方法、请求参数和响应结果。为保护用户隐私，@Log注解自动过滤参数中的password、oldPassword、newPassword、confirmPassword等敏感字段，以星号替代。日志数据写入sys_oper_log表持久化存储。系统还提供了登录日志的记录功能，通过SysLogininforService记录每次成功或失败的登录尝试，包含登录用户、IP地址、登录时间、操作系统和浏览器等信息。")

h2(doc, "4.8 前端实现")
body(doc, "前端基于Vue 2和Element UI构建，采用基于RuoYi-Vue的二次开发模式，在此基础上扩展了教学成果管理相关的页面和组件。前端通过统一的请求封装（utils/request.js）向后端发送HTTP请求，请求拦截器会自动添加Authorization头，响应拦截器会处理令牌过期等异常状态，弹出登录对话框引导用户重新登录。")
body(doc, "登录页面（views/login.vue）提供用户名密码登录和用户注册功能，支持图形验证码。登录成功后JWT令牌被存入浏览器的localStorage，后续所有请求自动附带该令牌。首页（views/dashboard）展示系统概览信息和快捷入口。成果管理页面（views/achievement）包括教师的成果列表页、成果编辑页和成果提交页，以及管理员的全校成果管理页和审核操作页。前端使用Element UI的Table组件实现数据列表展示和分页，Dialog组件实现弹窗编辑，Upload组件实现文件上传，Tag组件实现状态标签展示。")
body(doc, "权限控制在前端通过Vue指令（directive/permission）实现，在DOM元素渲染时根据用户权限集合决定是否显示。路由守卫（permission.js）在页面加载前检查用户登录状态和页面权限，未登录用户自动跳转登录页。前端开发服务器配置了代理规则，/dev-api前缀的请求会被代理到http://localhost:8080，并自动剥离/dev-api前缀，以匹配网关的路由规则。项目采用多环境配置（.env.development、.env.production、.env.staging），支持不同环境的快速切换。编译生产版本使用npm run build:prod命令，输出到edu-admin目录。")

print("Chapter 4 done. Writing Ch5...")
doc.save(OUTPUT)

# ======== CHAPTER 5 ========
h1(doc, "第五章 测试")

h2(doc, "5.1 测试环境")
body(doc, "系统测试在开发环境中进行，硬件环境为一台Windows 11桌面主机，配备32GB内存和1TB SSD存储。数据库采用MySQL 8.0.33，缓存采用Redis 5.0，注册中心采用Nacos 2.2.3单机模式，JDK版本为1.8.0。测试时所有微服务均在同一台机器上以默认端口运行，前端通过http://localhost访问。测试采用以功能验证为主的方法，从用户角度对系统的核心业务流程进行逐项验证，测试用例涵盖了成果申报与管理、审核流程、门户新闻、统计分析和用户权限五个核心模块。测试数据的准备通过执行edu_system.sql初始化脚本完成，该脚本包含了完整的业务表结构和多组示例数据。")

h2(doc, "5.2 功能测试")
body(doc, "本节以场景测试的形式对系统各核心功能模块进行验证，测试用例及其执行结果如下所述。")
body(doc, "测试用例1：教师新建成果并提交审核。前置条件为教师已登录系统，身份为普通教师角色。测试步骤：进入成果管理页面，点击新建按钮，填写成果标题为\"基于项目驱动的软件工程课程教学改革\"，选择成果类型为\"教学改革\"，申报等级为\"一等奖\"，在富文本编辑器中编写成果内容，上传PDF格式的证明材料附件，点击提交按钮。预期结果：系统提示提交成功，成果列表中显示新增的成果记录，状态为审核中。实际结果：与预期一致，成果成功提交并进入审核状态。")
body(doc, "测试用例2：教师编辑草稿成果。前置条件为存在一个草稿状态（0）的成果，当前登录用户为该成果的申报教师。测试步骤：在成果列表中点击编辑按钮，修改成果标题为\"基于项目驱动的软件工程课程教学改革与实践\"，保存。预期结果：成果信息更新成功，列表显示更新后的标题。实际结果：与预期一致，草稿成果可正常编辑。")
body(doc, "测试用例3：审核员审核通过成果。前置条件为存在一个审核中状态（2）的成果，审核员已登录。测试步骤：在审核页面查看成果详情，确认内容和附件无误，填写审核意见\"成果创新性强，教学效果显著，同意通过审核\"，点击通过按钮。预期结果：成果状态变为已通过（3），审核记录中新增一条通过记录。实际结果：与预期一致，审核操作正常执行，审核记录正确生成。")
body(doc, "测试用例4：审核员驳回成果。前置条件为存在一个审核中状态（2）的成果。测试步骤：在审核页面填写审核意见\"教学效果的实证数据不够充分，建议补充对比分析后重新提交\"，点击驳回按钮。预期结果：成果状态变为已驳回（4），审核记录中新增一条驳回记录，且教师可以查看到审核意见。实际结果：与预期一致。")
body(doc, "测试用例5：教师重新提交被驳回的成果。前置条件为存在一个已驳回状态（4）的成果，当前登录用户为该成果的申报教师。测试步骤：查看审核意见，根据意见修改成果内容，点击重新提交按钮。预期结果：成果状态变为审核中（2），再次进入审核队列。实际结果：与预期一致，驳回后的成果可正常重新提交审核。")
body(doc, "测试用例6：教师撤回审核中的成果。前置条件为存在一个审核中状态（2）的成果。测试步骤：在成果列表中点击撤回按钮。预期结果：成果状态变为草稿（0），审核记录保持不变。实际结果：与预期一致，审核中的成果可正常撤回为草稿。")
body(doc, "测试用例7：已通过成果禁止修改。前置条件为存在一个已通过状态（3）的成果，当前登录用户为该成果的申报教师。测试步骤：尝试编辑该成果的内容。预期结果：系统拒绝操作，提示\"该成果已通过审核，不可修改\"。实际结果：与预期一致，系统正确拦截了对已通过成果的修改操作。")
body(doc, "测试用例8：教师越权修改他人成果。前置条件为当前登录教师非成果所有者。测试步骤：尝试通过接口修改他人的成果。预期结果：系统拒绝操作，提示\"你无权修改他人的成果\"。实际结果：与预期一致，越权操作被正确拦截。")
body(doc, "测试用例9：门户新闻公开访问。前置条件为用户未登录，直接访问门户页面。测试步骤：在浏览器中打开门户新闻列表页，查看新闻列表，点击某个新闻进入详情页。预期结果：页面正常显示新闻列表和详情，无需登录，浏览量自动增加。实际结果：与预期一致，门户内容可正常公开访问。")
body(doc, "测试用例10：管理员查看统计数据。前置条件为管理员已登录。测试步骤：在首页查看成果总数、今日新增、本周新增、待审核数量等关键指标，分别进入按状态、类型、学院的统计页面查看详细数据。预期结果：各维度统计数据正确显示，与数据库实际数据一致。实际结果：与预期一致，统计计算结果正确。")
body(doc, "测试用例11：Excel导出功能。前置条件为管理员已登录。测试步骤：在成果列表页设置筛选条件，点击导出按钮，选择时间范围后执行导出。预期结果：下载一个Excel文件，包含符合筛选条件的成果数据，字段包括成果编号、标题、教师姓名、学院、状态、类型等。实际结果：与预期一致，Excel导出功能正常。")

h2(doc, "5.3 测试结果")
body(doc, "通过上述场景测试，系统各核心功能模块均通过了功能验证。成果的增删改查操作正常，状态流转逻辑严谨，权限校验和数据隔离机制有效。在测试过程中，系统对非法操作均能正确拦截并给出提示，包括越权修改他人成果、修改已通过成果、非法状态转换等场景。审核记录完整可追溯，在审核记录列表中可以查看任一成果的完整审核历史，包括通过和驳回的全部记录。门户新闻公开访问正常，无需登录即可浏览。统计数据计算准确，各维度相互一致。测试结果表明系统的核心业务功能已达到设计预期，可以满足教学成果管理的基本业务需求。")

print("Chapter 5 done. Writing Ch6 + End...")
doc.save(OUTPUT)

# ======== CHAPTER 6 + ENDING ========
h1(doc, "结束语")
body(doc, "本课题以郑州轻工业大学教学成果管理的实际需求为背景，设计并实现了一套基于Spring Cloud Alibaba微服务架构的教学成果管理系统。系统由网关服务、认证中心、系统管理服务、成果管理服务、文件服务和定时任务服务六个微服务组成，前端基于Vue 2和Element UI构建，数据层采用MySQL、Redis和文件对象存储。本文完成的主要工作包括：")
body(doc, "（1）分析了高校教学成果管理的业务需求，明确了教师、审核员和管理员三类角色的功能范围和权限边界，划分了成果管理、审核流程、门户新闻、统计分析和用户权限管理五个核心功能模块。")
body(doc, "（2）设计了基于微服务架构的系统整体方案，包括网关层、业务服务层和数据层的分层架构。完成了数据库E-R模型设计和主要业务表的结构定义，并制定了JWT+Redis双令牌认证、RBAC权限控制和五级数据范围过滤等安全方案。")
body(doc, "（3）实现了系统的各个核心模块，重点阐述了网关认证过滤器的工作原理、成果管理控制器的状态控制逻辑、审核流程的数据一致性保障以及前端页面的实现方式。通过场景测试验证了系统各模块的正确性和可用性。")
body(doc, "本系统实现了教学成果从在线申报、审核流转到门户公示的全流程数字化管理，可以为高校教学成果的统一管理和数据分析提供有效支持。系统也存在一些可以改进的方面，如当前审核流程为单级校级审核，后续可扩展为院级初审加校级终审的两级模式；AI查重接口已预留但尚未实现，可对接外部查重服务引入智能分析能力；前端界面可进一步适配移动端展示；微服务运维可引入分布式链路追踪和容器化部署等方案。")

print("Ending done. Adding references...")
doc.save(OUTPUT)

# ======== REFERENCES ========
h1(doc, "参考文献")

refs = [
    "[1] James Lewis, Martin Fowler. Microservices: a definition of this new architectural term[EB/OL]. (2014-03-25). https://martinfowler.com/articles/microservices.html.",
    "[2] Chris Richardson. Microservices Patterns: With examples in Java[M]. Shelter Island: Manning Publications, 2018: 1-50.",
    "[3] Sam Newman. Building Microservices: Designing Fine-Grained Systems[M]. Sebastopol: O'Reilly Media, 2015: 1-80.",
    "[4] 黄建平. 基于JWT与Redis的单点登录系统设计与实现[J]. 计算机技术与发展, 2020, 30(3): 25-31.",
    "[5] 孙卫琴. 精通Spring Cloud Alibaba——微服务架构与实战[M]. 北京: 北京大学出版社, 2020: 1-120.",
    "[6] 王志刚, 孙晓华. Spring Cloud微服务实战[M]. 北京: 人民邮电出版社, 2019: 45-102.",
    "[7] Spring Cloud Gateway Documentation[EB/OL]. https://docs.spring.io/spring-cloud-gateway/docs/current/reference/html/.",
    "[8] 崔永志, 贾喜平. Nacos实战与源码分析——微服务注册与配置管理[M]. 北京: 电子工业出版社, 2021: 23-89.",
    "[9] 刘斌, 陈建峰. MyBatis从入门到精通[M]. 北京: 清华大学出版社, 2017: 78-135.",
    "[10] Redis官方文档[EB/OL]. https://redis.io/docs/latest/.",
    "[11] Vue.js v2官方指南[EB/OL]. https://v2.vuejs.org/v2/guide/.",
    "[12] 教育部. 关于深化本科教育教学改革全面提高人才培养质量的意见[EB/OL]. (2019-10-08). http://www.moe.gov.cn/srcsite/A08/s7056/201910/t20191011_402759.html.",
    "[13] Roy Thomas Fielding. Architectural Styles and the Design of Network-based Software Architectures[D]. Irvine: University of California, 2000.",
    "[14] Cesare Pautasso, Olaf Zimmermann, Frank Leymann. RESTful Web Services vs. Big Web Services: Making the Right Architectural Decision[C]. Beijing: Proceedings of the 17th International World Wide Web Conference (WWW 2008), 2008: 805-814.",
    "[15] 耿祥义, 张跃平. Java并发编程实战[M]. 北京: 机械工业出版社, 2015: 120-156.",
    "[16] Irakli Nadareishvili, Ronnie Mitra, Matt McLarty, Mike Amundsen. Microservice Architecture: Aligning Principles, Practices, and Culture[M]. Sebastopol: O'Reilly Media, 2016: 1-60.",
    "[17] Martin Fowler. Patterns of Enterprise Application Architecture[M]. Boston: Addison-Wesley, 2002: 1-100.",
    "[18] Eric Evans. Domain-Driven Design: Tackling Complexity in the Heart of Software[M]. Boston: Addison-Wesley, 2003: 1-80.",
    "[19] M. Jones, J. Bradley, N. Sakimura. JSON Web Token (JWT): RFC 7519[S]. IETF, 2015. https://tools.ietf.org/html/rfc7519.",
    "[20] 李俊杰. 基于RBAC模型的权限管理系统设计与实现[D]. 成都: 电子科技大学, 2018.",
]
for rt in refs:
    ref_para(doc, rt)

# ======== THANKS ========
h1(doc, "致谢")
body(doc, "在本文完成之际，首先向我的指导教师致以最诚挚的感谢。在毕业设计的整个过程中，指导教师从选题确定、技术方案论证、系统设计到论文撰写的每个环节都给予了细致的指导和建设性意见，帮助我克服了诸多技术难题和思路困惑。")
body(doc, "感谢软件学院各位任课教师在四年本科学习中的辛勤教导，为我奠定了软件开发的理论基础和实践能力。感谢实验室同学们在项目开发过程中的互助合作，大家的讨论和交流让我对微服务架构有了更深入的理解。")
body(doc, "感谢郑州轻工业大学提供了优越的学习环境和实验条件，让我能够将所学知识应用于实际项目中。最后感谢家人在求学期间的理解和鼓励，他们的无私支持是我完成学业的重要保障。")

# ======== SAVE FINAL ========
doc.save(OUTPUT)
print("SAVED. Inserting images...")

# ======== INSERT IMAGES ========
doc = Document(OUTPUT)
fig_map = {}
for fn in os.listdir(FIGDIR):
    if fn.endswith('.png'):
        m = re.match(r'\u56fe(\d+)-(\d+)', fn)
        if m: fig_map[(m.group(1), m.group(2))] = os.path.join(FIGDIR, fn)

inserted = 0
fig_caps = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'\u56fe(\d+)-(\d+)', t)
    if m and p.runs and p.runs[0].bold and len(t) < 30:
        fig_caps.append((i, m.group(1), m.group(2), t))

for para_idx, ch, fnum, caption in reversed(fig_caps):
    key = (ch, fnum)
    if key not in fig_map: continue
    png_path = fig_map[key]
    cap_elem = doc.paragraphs[para_idx]._element
    parent = cap_elem.getparent()
    cap_pos = list(parent).index(cap_elem)
    
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Cm(0)
    img_para.paragraph_format.space_after = Cm(0)
    
    try:
        img_para.add_run().add_picture(png_path, width=Inches(4.8))
    except: continue
    
    img_elem = img_para._element
    img_elem.getparent().remove(img_elem)
    parent.insert(cap_pos, img_elem)
    inserted += 1

doc.save(OUTPUT)
print(f"COMPLETE: {inserted} images inserted, {len(refs)} references, {len(doc.tables)} tables")
print(f"File: {OUTPUT}")
