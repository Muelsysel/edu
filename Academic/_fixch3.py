# -*- coding: utf-8 -*-
"""Add Ch3 design diagrams to thesis-final.docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, re

OUTPUT = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\thesis-final.docx"
FIGDIR = r"D:\.develop\edu-achievement01\edu-achievement-master\Academic\figures"
doc = Document(OUTPUT)

def fig_cap_ins(parent, pos, text):
    wp = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); wp.insert(0,pPr)
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"),"center"); pPr.append(jc)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:line"),"360"); sp.set(qn("w:lineRule"),"auto")
    pPr.append(sp)
    wr = OxmlElement("w:r"); wp.append(wr)
    wrPr = OxmlElement("w:rPr"); wr.append(wrPr)
    wrSz = OxmlElement("w:sz"); wrSz.set(qn("w:val"),"21"); wrPr.append(wrSz)
    wrB = OxmlElement("w:b"); wrPr.append(wrB)
    wrRf = OxmlElement("w:rFonts")
    wrRf.set(qn("w:eastAsia"),"宋体"); wrRf.set(qn("w:ascii"),"Times New Roman")
    wrPr.append(wrRf)
    wrt = OxmlElement("w:t"); wrt.text = text; wr.append(wrt)
    parent.insert(pos, wp)
    return pos + 1

def body_ins(parent, pos, text):
    wp = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr"); wp.insert(0,pPr)
    ind = OxmlElement("w:ind"); ind.set(qn("w:firstLine"),"480"); pPr.append(ind)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:line"),"360"); sp.set(qn("w:lineRule"),"auto")
    pPr.append(sp)
    wr = OxmlElement("w:r"); wp.append(wr)
    wrPr = OxmlElement("w:rPr"); wr.append(wrPr)
    wrSz = OxmlElement("w:sz"); wrSz.set(qn("w:val"),"24"); wrPr.append(wrSz)
    wrRf = OxmlElement("w:rFonts")
    wrRf.set(qn("w:eastAsia"),"宋体"); wrRf.set(qn("w:ascii"),"Times New Roman")
    wrPr.append(wrRf)
    wrt = OxmlElement("w:t"); wrt.text = text; wr.append(wrt)
    parent.insert(pos, wp)
    return pos + 1

# Find 3.4
sec34 = None
for i, p in enumerate(doc.paragraphs):
    if "3.4" in p.text and "安全" in p.text and len(p.text.strip()) < 20:
        sec34 = i; break

if sec34:
    ref = doc.paragraphs[sec34]._element
    parent = ref.getparent()
    pos = list(parent).index(ref)

    # 图3-4: 登录认证时序图
    b1 = "认证流程是系统安全体系的第一道防线。网关层的AuthFilter全局过滤器对每个请求进行令牌验证，合法的请求被添加用户身份Header后转发到下游服务。整个登录认证的完整交互过程如图3-4所示。"
    c1 = "图3-4 登录认证时序图"
    pos = body_ins(parent, pos, b1)
    pos = fig_cap_ins(parent, pos, c1)

    # 图3-5: 成果状态控制时序图
    b2 = "教师修改成果时的状态控制涉及三层校验：身份确认、状态合法性判断和数据库动态更新。首先从数据库加载旧数据进行对比，然后依次检查是否为成果所有者以及当前状态是否允许编辑，最终使用MyBatis动态SQL仅更新非空字段。具体流程如图3-5所示。"
    c2 = "图3-5 成果状态控制时序图"
    pos = body_ins(parent, pos, b2)
    pos = fig_cap_ins(parent, pos, c2)

    # 图3-6: 审核事务流程时序图
    b3 = "审核操作是两个数据库操作的事务组合——同时更新成果主表的status字段和向审核记录表插入新记录。为保证原子性，系统通过@Transactional注解将整个审核流程包裹在同一事务中，具体过程如图3-6所示。"
    c3 = "图3-6 审核事务流程时序图"
    pos = body_ins(parent, pos, b3)
    pos = fig_cap_ins(parent, pos, c3)

    # 图3-7: 安全与权限控制核心类图
    b4 = "系统的安全体系由多个切面和拦截器协同实现。接口权限校验通过PreAuthorizeAspect切面在方法执行前验证权限注解，数据权限通过DataScopeAspect切面在MyBatis查询时动态注入过滤条件。核心安全组件的静态关系如图3-7所示。"
    c4 = "图3-7 安全与权限控制核心类图"
    pos = body_ins(parent, pos, b4)
    pos = fig_cap_ins(parent, pos, c4)

    # Insert PNGs
    png_map = {}
    for fn in os.listdir(FIGDIR):
        if fn.endswith(".png"):
            m = re.match(r"图(\d+)-(\d+)", fn)
            if m: png_map[f"{m.group(1)}-{m.group(2)}"] = os.path.join(FIGDIR, fn)

    for key in ["4-1", "4-2", "4-3", "4-4"]:
        if key in png_map:
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                img_p.add_run().add_picture(png_map[key], width=Inches(4.8))
                img_e = img_p._element
                img_e.getparent().remove(img_e)
                parent.insert(pos, img_e)
                pos += 1
                print(f"PNG {key} inserted")
            except Exception as e:
                print(f"PNG {key} err: {e}")

print(f"Ch3 design diagrams done. Paras: {len(doc.paragraphs)}")
doc.save(OUTPUT)
print("Saved.")
